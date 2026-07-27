"""Gera a documentacao tecnica (DOCX) das planilhas e informacoes que
compoem o consolidado_equalizacao.xlsx, detalhando cada coluna, seu
conteudo e as regras de negocio envolvidas.

Uso:
    python scripts/gerar_documentacao_planilhas.py

Gera docs/Documentacao_Planilha_Consolidada.docx.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent.parent
CAMINHO_SAIDA = BASE_DIR / "docs" / "Documentacao_Planilha_Consolidada.docx"

A4_LARGURA = Cm(21)
A4_ALTURA = Cm(29.7)
COR_TITULO = RGBColor(0x1F, 0x3A, 0x5F)
COR_CODIGO = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# Utilitarios de formatacao (mesmo padrao visual do relatorio gerencial)
# ---------------------------------------------------------------------------

def _config_pagina_a4(document: Document) -> None:
    for section in document.sections:
        section.page_width = A4_LARGURA
        section.page_height = A4_ALTURA
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def _marcar_linha_cabecalho_repetido(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _adicionar_campo_pagina(paragraph, texto_antes: str = "", texto_depois: str = "") -> None:
    paragraph.add_run(texto_antes)

    def _campo(tipo: str):
        r = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {tipo} "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_end)

    _campo("PAGE")
    paragraph.add_run(" de ")
    _campo("NUMPAGES")
    paragraph.add_run(texto_depois)


def _rodape(document: Document, data_geracao: datetime) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        _adicionar_campo_pagina(paragraph, texto_antes="Pagina ",
                                 texto_depois=f" | Documentacao gerada em {data_geracao.strftime('%d/%m/%Y %H:%M')}")


def _titulo(document: Document, texto: str, nivel: int = 1):
    h = document.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = COR_TITULO
    return h


def _paragrafo(document: Document, texto: str, italico: bool = False):
    p = document.add_paragraph()
    run = p.add_run(texto)
    run.italic = italico
    return p


def _tabela_colunas(document: Document, linhas: list[tuple], larguras_cm: tuple = (4.2, 2.3, 5.5, 5.0)) -> None:
    """Monta a tabela padrao de documentacao de colunas: Coluna | Tipo |
    Conteudo | Regra de negocio."""
    tabela = document.add_table(rows=1, cols=4)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Light Grid Accent 1"

    cabecalhos = ["Coluna", "Tipo", "Conteudo", "Regra de negocio / observacoes"]
    for i, texto in enumerate(cabecalhos):
        celula = tabela.rows[0].cells[i]
        celula.text = texto
        for p in celula.paragraphs:
            for r in p.runs:
                r.bold = True
        celula.width = Cm(larguras_cm[i])
    _marcar_linha_cabecalho_repetido(tabela.rows[0])

    for coluna, tipo, conteudo, regra in linhas:
        celulas = tabela.add_row().cells
        celulas[0].text = coluna
        for p in celulas[0].paragraphs:
            for r in p.runs:
                r.font.name = "Consolas"
                r.font.color.rgb = COR_CODIGO
        celulas[1].text = tipo
        celulas[2].text = conteudo
        celulas[3].text = regra
        for i, largura in enumerate(larguras_cm):
            celulas[i].width = Cm(largura)
    document.add_paragraph()


def _aba(document: Document, numero: str, nome_aba: str, arquivo_origem: str, descricao: str,
         granularidade: str, colunas: list[tuple]) -> None:
    _titulo(document, f"{numero}. Aba `{nome_aba}`", nivel=2)
    _paragrafo(document, descricao)
    p = document.add_paragraph()
    p.add_run("Granularidade: ").bold = True
    p.add_run(granularidade)
    p2 = document.add_paragraph()
    p2.add_run("Tambem disponivel em: ").bold = True
    p2.add_run(arquivo_origem)
    _tabela_colunas(document, colunas)


# ---------------------------------------------------------------------------
# Conteudo
# ---------------------------------------------------------------------------

TEXTO_INTRODUCAO = """
Este documento descreve, coluna a coluna, o conteudo e as regras de negocio de cada aba do arquivo output/consolidado_equalizacao.xlsx, gerado automaticamente pelo Sistema de Equalizacao TRT12 a cada processamento (python main.py, python main.py --atualizar, ou ao abrir/atualizar o dashboard Streamlit).

O consolidado e a base completa da equalizacao, sem qualquer filtro aplicado: reune, em um unico arquivo Excel com oito abas, as movimentacoes processuais consolidadas, o resumo analitico por processo, a auditoria de episodios de passagem por triagem, os cenarios comparativos com e sem equalizacao, os fluxos entre unidades, o catalogo de anomalias de qualidade e a relacao de unidades citadas nos arquivos de entrada que nao constam no arquivo de classificacao.

As mesmas informacoes tambem sao exportadas em CSV/Parquet individuais na pasta output/ (movimentacoes_consolidadas, processos_consolidados, episodios_equalizacao, comparativo_unidades, fluxos_equalizacao, anomalias, unidades_nao_classificadas) -- o Excel consolidado reune tudo em um so lugar, com cabecalho fixo e autofiltro em cada aba, para facilitar a analise fora do dashboard.
""".strip()

TEXTO_CONVENCOES = """
Colunas `_norm` vs. `_original`: para nomes de unidade, indicador, classe judicial e municipio, o sistema mantem duas versoes de cada valor. A coluna sem sufixo (ou com sufixo `_original`) preserva o texto exatamente como veio do CSV de origem, para exibicao. A coluna `_norm` traz a versao normalizada -- maiusculas, sem acento, espacos duplicados colapsados, ordinais unificados (1a, 1A, 1o viram 1), hifens e pontuacao normalizados -- usada internamente para comparar, agrupar e casar registros. O sistema nunca faz correspondencia aproximada (fuzzy matching): duas grafias que normalizam para textos diferentes permanecem diferentes.

Convencao de contagem: "processo" e sempre `COUNT DISTINCT` do numero do processo; "movimentacao" conta linhas de movimentacao; "episodio" conta episodios de passagem por triagem (podem existir varios por processo); "unidade" conta unidades permanentes (as classificadas como CEDENTE, NEUTRA ou DESTINATARIA no arquivo de classificacao -- unidades de triagem nunca sao contadas como unidade permanente).

Datas: todas as colunas de data sao armazenadas como data/hora (sem fuso horario, pois o Excel nao suporta timezone). Quando a coluna Data do CSV de origem trouxer apenas a data, a hora fica zerada (00:00:00).

Celulas vazias: uma unidade vazia ("") em uma coluna de texto indica que o campo nao pode ser determinado a partir do historico do processo (por exemplo, um processo sem nenhuma distribuicao valida nao tem `unidade_distribuicao_inicial`). Uma celula vazia em uma coluna de data (`NaT`) tem o mesmo significado.

Ordem de reconstrucao: o historico cronologico completo de cada processo e reconstruido, e a deteccao de equalizacao e calculada, ANTES de qualquer filtro de periodo ou unidade. Isso garante que uma equalizacao nao deixe de ser identificada so porque a distribuicao inicial ocorreu antes do periodo selecionado no dashboard ou o destino ocorreu depois dele -- os filtros do dashboard atuam sobre esta base ja pronta, nunca sobre o historico bruto.

Movimentacoes RESTITUIDO: movimentacoes cujo indicador seja "Restituido Por Redistribuicao" sao classificadas como RESTITUIDO e permanecem integralmente na aba Movimentacoes (para rastreabilidade), mas sao excluidas de toda e qualquer analise -- nao contam em quantidade_movimentos nem quantidade_redistribuicoes, nao participam da reconstrucao de historico usada para detectar equalizacao, nao geram anomalia de ordem cronologica ambigua e nao entram no relatorio de unidades nao classificadas. Se todas as movimentacoes de um processo forem RESTITUIDO, esse processo nao aparece na aba Processos (nao ha nada a analisar), mas seus registros continuam na aba Movimentacoes.
""".strip()

TEXTO_REGRA_EQUALIZACAO = """
Uma equalizacao e detectada quando o historico cronologico do processo contem a sequencia UNIDADE PERMANENTE -> UNIDADE DE TRIAGEM -> UNIDADE PERMANENTE. O algoritmo percorre o historico ja ordenado do processo mantendo o registro da ultima unidade permanente vista; ao encontrar uma unidade de triagem, abre um episodio; ao encontrar a proxima unidade permanente, fecha o episodio. Isso cobre naturalmente processos com mais de uma passagem pela triagem (multiplos episodios).

Uma unidade e considerada "de triagem" quando seu nome normalizado contem a palavra TRIAGEM (comparacao insensivel a maiusculas/minusculas e acentos). Movimentacoes em unidades que nao sao nem permanentes (constam no arquivo de classificacao) nem de triagem sao ignoradas na deteccao de episodios -- ja ficam registradas como anomalia UNIDADE_NAO_CLASSIFICADA em outra etapa e nao servem como ancora confiavel de origem/destino.

Cada episodio e classificado em um dos tres tipos:
- VALIDA: a unidade anterior a triagem esta classificada como CEDENTE, a primeira unidade permanente posterior a triagem esta classificada como DESTINATARIA, e os dois movimentos de entrada (na triagem e no destino) sao redistribuicoes.
- FORA_PADRAO: existe uma unidade permanente definitiva apos a triagem, mas alguma das condicoes acima nao e atendida (origem nao e cedente, destino nao e destinataria, ou algum dos dois movimentos de entrada nao e redistribuicao).
- INCOMPLETA: nao ha nenhuma unidade permanente definitiva apos a entrada na triagem (o processo permanece la, ou o historico termina antes de sair).

Quando um processo tem mais de um episodio, todos vao para a aba Episodios_Equalizacao, mas apenas o primeiro episodio VALIDO define a origem/destino "oficiais" do processo (colunas unidade_cedente_equalizacao / unidade_destinataria_equalizacao na aba Processos); se nao houver nenhum episodio valido, usa-se o primeiro episodio de qualquer tipo apenas para fins informativos (motivo_invalidade_equalizacao). O fato de haver mais de um episodio e sinalizado como anomalia MULTIPLOS_EPISODIOS_EQUALIZACAO, mas o processo continua contado uma unica vez em todos os indicadores.
""".strip()

TEXTO_REGRA_CENARIOS = """
A aba Comparativo_Unidades apresenta dois cenarios, sempre com a mesma quantidade total de processos distintos (cada processo e contado exatamente uma vez em cada cenario):

- COM_EQUALIZACAO: cada processo e contado na sua unidade_final_real (ultima unidade permanente do historico cronologico), isto e, onde o processo de fato esta hoje.
- SEM_EQUALIZACAO: cenario simulado. Processos com equalizacao_valida = verdadeiro sao recontados na unidade_cedente_equalizacao (a unidade cedente imediatamente anterior a triagem) -- ou seja, "como se o Sistema de Equalizacao nao existisse, o processo teria ficado na origem". Os demais processos (sem equalizacao valida) permanecem na sua unidade_final_real.

Ambos os cenarios incluem todas as unidades permanentes do arquivo de classificacao, mesmo as que nao receberam nenhum processo (quantidade = 0), para que as medidas de dispersao (media, MAD, desvio-padrao, coeficiente de variacao) considerem o universo completo de unidades.

Processos cuja unidade relevante (final real ou simulada) nao pertence ao universo de unidades classificadas nao aparecem nas linhas da aba Comparativo_Unidades (ela cobre apenas unidades permanentes classificadas), mas continuam contados normalmente em "casos novos" -- a verificacao de consistencia dos totais leva isso em conta.
""".strip()

TEXTO_REGRA_MAGISTRADO = """
Alem das medidas por unidade (cada unidade com o mesmo peso), a aba Comparativo_Unidades traz um segundo grupo de medidas, ponderadas pela quantidade de magistrados de cada unidade -- util para comparar a carga de trabalho real dos magistrados, e nao apenas a quantidade bruta de processos por unidade.

Para cada unidade, processos_por_magistrado = quantidade de processos da unidade / quantidade de magistrados da unidade. O centro de referencia desse grupo NAO e a media simples dessas taxas por unidade; e a media ponderada do sistema: soma dos processos de todas as unidades com quantidade de magistrados valida, dividida pela soma dos magistrados dessas mesmas unidades. Isso equivale ao "total de processos por magistrado" do sistema como um todo, e evita que unidades pequenas (com taxas mais instaveis) distorcam a media tanto quanto distorceriam em uma media simples.

As medidas de dispersao deste grupo (MAD, desvio-padrao, coeficiente de variacao, amplitude, unidades acima/abaixo da media) sao calculadas da mesma forma que as medidas por unidade, mas em torno dessa media ponderada -- ou seja, medem o quanto a carga por magistrado de cada unidade se afasta do patamar medio do sistema, nao da media das proprias taxas.

Como o total de processos e o total de magistrados nao mudam entre os cenarios COM_EQUALIZACAO e SEM_EQUALIZACAO (a equalizacao apenas redistribui em qual unidade cada processo e contado), a media ponderada por magistrado e sempre a MESMA nos dois cenarios -- o que muda, e o que evidencia o efeito do Sistema de Equalizacao, e a dispersao em torno dela (normalmente menor no cenario com equalizacao).

Unidades sem quantidade de magistrados valida (arquivo de classificacao sem essa coluna, ou valor ausente/invalido para a unidade) ficam de fora do calculo da media ponderada e das medidas de dispersao deste grupo -- processos_por_magistrado e as diferencas ficam vazios (N/A) para essas unidades, mas elas continuam normalmente nas demais colunas e no grupo de medidas por unidade.
""".strip()


def _colunas_resumo():
    return [
        ("indicador", "Texto", "Nome do indicador de qualidade ou metadado da execucao.",
         "Lista fixa: registros lidos, quantidade de arquivos, duplicidades removidas, registros validos, "
         "registros com anomalia, registros RESTITUIDO, processos distintos, unidades nao classificadas, "
         "data minima/maxima, data de geracao, arquivos utilizados e avisos (um por linha)."),
        ("valor", "Texto / numero / data", "Valor correspondente ao indicador da mesma linha.",
         "Sem regra de calculo propria; reflete o valor computado pelas demais abas na execucao que gerou o arquivo."),
    ]


def _colunas_processos():
    return [
        ("processo", "Texto", "Numero do processo, como aparece no CSV de origem (zeros a esquerda e pontuacao "
         "CNJ preservados).", "Nunca convertido para numero, para nao perder zeros a esquerda nem alterar a "
         "identificacao processual. Chave de exibicao do processo."),
        ("processo_norm", "Texto", "Numero do processo normalizado (maiusculas, sem espacos extras).",
         "Chave interna de casamento/agrupamento entre todas as abas e com o arquivo de log usado por "
         "comparar_processos.py."),
        ("data_caso_novo", "Data/hora", "Data da primeira movimentacao com indicador DISTRIBUICAO do processo.",
         "Vazio quando o processo nao possui nenhuma distribuicao valida no historico analisado (ver anomalia "
         "PRIMEIRO_MOVIMENTO_NAO_DISTRIBUICAO). E a base temporal padrao do filtro de periodo do dashboard."),
        ("unidade_distribuicao_inicial", "Texto (normalizado)", "Unidade da primeira distribuicao do processo, "
         "normalizada.", "Extraida da mesma movimentacao usada para data_caso_novo."),
        ("unidade_distribuicao_inicial_original", "Texto", "Mesma unidade acima, com a grafia original do CSV.",
         "Uso apenas de exibicao."),
        ("unidade_final_real", "Texto (normalizado)", "Ultima unidade PERMANENTE (nao triagem) registrada no "
         "historico do processo.", "Define o cenario COM_EQUALIZACAO no Comparativo_Unidades. Se o processo "
         "nunca passou por uma unidade permanente, fica vazio."),
        ("unidade_final_real_original", "Texto", "Mesma unidade acima, com a grafia original do CSV.",
         "Uso apenas de exibicao."),
        ("data_ultimo_movimento", "Data/hora", "Data da ultima movimentacao (analisada) do processo.",
         "Uma das quatro dimensoes de data selecionaveis como base temporal do filtro de periodo."),
        ("classe_judicial", "Texto (normalizado)", "Classe judicial da ultima movimentacao do processo.",
         "Quando movimentos do mesmo processo tem classes diferentes, isso e sinalizado como anomalia "
         "CLASSE_JUDICIAL_DIVERGENTE, mas esta coluna sempre reflete a classe mais recente."),
        ("quantidade_movimentos", "Numero inteiro", "Quantidade de movimentacoes do processo consideradas na "
         "analise.", "NAO inclui movimentacoes com indicador RESTITUIDO (desconsideradas de toda analise)."),
        ("quantidade_redistribuicoes", "Numero inteiro", "Quantidade de movimentacoes com indicador REDISTRIBUICAO.",
         "Idem: RESTITUIDO nao entra nesta contagem, mesmo contendo o texto 'redistribuicao' no indicador original."),
        ("passou_por_triagem", "Booleano", "Indica se alguma movimentacao (analisada) do processo ocorreu em "
         "unidade de triagem.", "Verdadeiro mesmo quando a passagem pela triagem nao configura uma equalizacao "
         "valida (ver equalizacao_detectada / equalizacao_valida)."),
        ("municipio_sede", "Texto (normalizado)", "Municipio sede da ultima movimentacao do processo.", "-"),
        ("municipio_sede_original", "Texto", "Mesmo municipio acima, com a grafia original do CSV.", "-"),
        ("arquivo_origem_ref", "Texto", "Nome do arquivo CSV da primeira movimentacao do processo.",
         "Usado apenas como referencia de rastreabilidade; o historico completo por arquivo/linha esta na aba "
         "Movimentacoes."),
        ("numero_linha_arquivo_ref", "Numero inteiro", "Numero da linha fisica, no arquivo acima, da primeira "
         "movimentacao do processo.", "Considera as linhas de resumo descartadas no inicio do CSV (a contagem "
         "reflete a linha real do arquivo, nao um indice relativo)."),
        ("equalizacao_detectada", "Booleano", "Verdadeiro quando existe ao menos um episodio de passagem pela "
         "triagem que chegou a uma unidade permanente de destino (tipo VALIDA ou FORA_PADRAO).",
         "Diferente de passou_por_triagem: um processo pode ter entrado na triagem e nunca saido (episodio "
         "INCOMPLETA), e nesse caso equalizacao_detectada e falso."),
        ("equalizacao_valida", "Booleano", "Verdadeiro quando existe ao menos um episodio classificado como VALIDA.",
         "Ver regra completa de deteccao de equalizacao mais adiante neste documento. E o criterio oficial para "
         "os indicadores gerenciais de 'processos equalizados'."),
        ("quantidade_episodios_triagem", "Numero inteiro", "Quantidade total de episodios de passagem pela "
         "triagem do processo (validos, fora de padrao e incompletos somados).",
         "Quando maior que 1, gera a anomalia MULTIPLOS_EPISODIOS_EQUALIZACAO."),
        ("unidade_cedente_equalizacao", "Texto (normalizado)", "Unidade imediatamente anterior a triagem no "
         "episodio 'principal' do processo (o primeiro episodio valido, ou o primeiro episodio de qualquer tipo "
         "se nao houver nenhum valido).", "Fica vazio quando o processo nao tem nenhum episodio de triagem. Usada "
         "para montar unidade_simulada_sem_equalizacao quando equalizacao_valida e verdadeiro."),
        ("unidade_cedente_equalizacao_original", "Texto", "Mesma unidade acima, com a grafia oficial do arquivo "
         "de classificacao.", "-"),
        ("unidade_destinataria_equalizacao", "Texto (normalizado)", "Primeira unidade permanente apos a triagem "
         "no episodio 'principal' do processo.", "Fica vazio quando o episodio principal e do tipo INCOMPLETA "
         "(nao ha destino)."),
        ("unidade_destinataria_equalizacao_original", "Texto", "Mesma unidade acima, com a grafia oficial do "
         "arquivo de classificacao.", "-"),
        ("data_entrada_triagem", "Data/hora", "Data em que o processo entrou na unidade de triagem, no episodio "
         "principal.", "Uma das quatro dimensoes de data do processo (nao selecionavel diretamente como base do "
         "filtro de periodo do dashboard, mas disponivel para consulta/analise)."),
        ("data_destino_equalizacao", "Data/hora", "Data em que o processo chegou a unidade de destino, no "
         "episodio principal.", "Vazio quando o episodio principal e INCOMPLETA. Selecionavel como base temporal "
         "'conclusao da equalizacao' no filtro de periodo do dashboard."),
        ("motivo_invalidade_equalizacao", "Texto", "Explicacao textual de por que o episodio principal nao e "
         "valido (quando equalizacao_valida e falso mas ha um episodio principal).",
         "Vazio quando equalizacao_valida e verdadeiro. Pode conter mais de um motivo, separados por ';' (ex.: "
         "origem nao e CEDENTE E movimento de entrada na triagem nao e redistribuicao)."),
        ("unidade_simulada_sem_equalizacao", "Texto (normalizado)", "Unidade onde o processo seria contado no "
         "cenario simulado 'sem equalizacao'.", "= unidade_cedente_equalizacao quando equalizacao_valida e "
         "verdadeiro; = unidade_final_real caso contrario (inclusive quando ha passagem por triagem invalida ou "
         "incompleta). Ver secao de regras dos cenarios comparativos."),
        ("unidade_simulada_sem_equalizacao_original", "Texto", "Mesma unidade acima, com a grafia original.", "-"),
        ("classificacao_unidade_final_real", "Texto", "Classificacao (CEDENTE/NEUTRA/DESTINATARIA) da unidade "
         "final real, segundo o arquivo de classificacao.", "NAO_CLASSIFICADA quando a unidade nao consta no "
         "arquivo de classificacao; vazio quando unidade_final_real e vazio."),
        ("classificacao_unidade_distribuicao_inicial", "Texto", "Classificacao da unidade de distribuicao "
         "inicial.", "Mesma regra de NAO_CLASSIFICADA/vazio acima."),
        ("classificacao_unidade_cedente_equalizacao", "Texto", "Classificacao da unidade cedente da "
         "equalizacao.", "Mesma regra de NAO_CLASSIFICADA/vazio acima."),
        ("classificacao_unidade_destinataria_equalizacao", "Texto", "Classificacao da unidade destinataria da "
         "equalizacao.", "Mesma regra de NAO_CLASSIFICADA/vazio acima."),
        ("classificacao_unidade_simulada_sem_equalizacao", "Texto", "Classificacao da unidade simulada sem "
         "equalizacao.", "Mesma regra de NAO_CLASSIFICADA/vazio acima."),
        ("tem_anomalia", "Booleano", "Indica se o processo tem pelo menos um registro na aba Anomalias.",
         "Usado pelo filtro 'Processos com anomalia' do dashboard."),
    ]


def _colunas_movimentacoes():
    return [
        ("processo", "Texto", "Numero do processo apos remocao de espacos externos (sem outra normalizacao).",
         "Preserva zeros a esquerda e pontuacao; nunca convertido para numero."),
        ("indicador", "Texto", "Texto do indicador tal como veio do CSV apos padronizacao de nome de coluna.",
         "Igual a indicador_original nesta aba (mantido por compatibilidade com o nome de coluna do CSV de "
         "entrada)."),
        ("municipio_sede", "Texto", "Municipio sede tal como veio do CSV.", "-"),
        ("orgao_julgador", "Texto", "Orgao julgador (unidade) tal como veio do CSV.", "-"),
        ("classe_judicial", "Texto", "Classe judicial tal como veio do CSV.", "-"),
        ("data", "Data/hora ou texto", "Valor original da coluna Data do CSV (pode incluir data e hora).",
         "Ver data_dt para a versao convertida e validada."),
        ("arquivo_origem", "Texto", "Nome do arquivo CSV de onde a movimentacao foi lida.",
         "Usado como criterio de desempate na ordenacao cronologica e em toda referencia de rastreabilidade "
         "(anomalias, episodios etc.)."),
        ("data_arquivo", "Data", "Data extraida do nome do arquivo (processos_DD-MM-YYYY.csv).",
         "Informacao auxiliar de rastreabilidade; a cronologia principal das movimentacoes usa a coluna Data, "
         "nao esta."),
        ("ordem_arquivo", "Numero inteiro", "Posicao do arquivo na ordem cronologica de processamento (0-based).",
         "Calculada ordenando os arquivos por data_arquivo e, em caso de empate, pelo nome do arquivo. Usada "
         "como criterio de desempate na ordenacao."),
        ("numero_linha_arquivo", "Numero inteiro", "Numero da linha fisica da movimentacao no arquivo CSV "
         "original (1-based).", "Considera as linhas de resumo do inicio do arquivo, que sao descartadas antes "
         "da leitura da tabela -- o numero reflete a linha real do arquivo."),
        ("data_processamento", "Data/hora", "Momento em que o arquivo foi lido pelo sistema.",
         "Metadado de auditoria da execucao; nao usado em nenhuma regra de negocio."),
        ("processo_original", "Texto", "Copia do numero do processo antes de qualquer normalizacao.",
         "Identica a coluna 'processo' nesta aba (mantida para rastreabilidade do pipeline interno)."),
        ("processo_norm", "Texto", "Numero do processo normalizado (maiusculas).",
         "Chave de agrupamento/casamento com as demais abas."),
        ("indicador_original", "Texto", "Texto do indicador exatamente como veio do CSV.", "-"),
        ("indicador_norm", "Texto", "Indicador classificado em uma categoria fixa.",
         "Valores possiveis: DISTRIBUICAO, REDISTRIBUICAO, RESTITUIDO, OUTRO, NAO_IDENTIFICADO. A busca por "
         "RESTITUIDO precede as demais (pois 'Restituido Por Redistribuicao' contem o trecho 'redistribuicao'); "
         "em seguida REDISTRIBUICAO precede DISTRIBUICAO pelo mesmo motivo. NAO_IDENTIFICADO quando o campo "
         "esta vazio."),
        ("orgao_julgador_original", "Texto", "Nome da unidade exatamente como veio do CSV.",
         "Usado para exibicao em tabelas e no relatorio DOCX."),
        ("orgao_norm", "Texto", "Nome da unidade normalizado.",
         "Chave usada para casar com o arquivo de classificacao e para identificar unidades de triagem "
         "(contem a palavra TRIAGEM)."),
        ("classe_judicial_original", "Texto", "Classe judicial exatamente como veio do CSV.", "-"),
        ("classe_norm", "Texto", "Classe judicial normalizada.", "-"),
        ("municipio_sede_original", "Texto", "Municipio sede exatamente como veio do CSV.", "-"),
        ("municipio_norm", "Texto", "Municipio sede normalizado.", "-"),
        ("data_dt", "Data/hora", "Data da movimentacao convertida para data/hora (formato brasileiro, dia "
         "primeiro).", "Vazia (NaT) quando a coluna Data original nao pode ser interpretada como data -- gera "
         "a anomalia DATA_INVALIDA. E a coluna usada para ordenar cronologicamente o historico."),
        ("data_valida", "Booleano", "Verdadeiro quando data_dt pode ser interpretada.",
         "Espelha diretamente se data_dt e ou nao NaT."),
        ("unidade_e_triagem", "Booleano", "Verdadeiro quando orgao_norm contem a palavra TRIAGEM.",
         "Comparacao insensivel a maiusculas/minusculas e acentos (feita sobre o nome ja normalizado)."),
        ("sequencia_movimento", "Numero inteiro", "Posicao da movimentacao na ordem cronologica do processo "
         "(1, 2, 3...).", "Ordenacao por: data (data_dt) -> data extraida do arquivo (data_arquivo) -> ordem do "
         "arquivo (ordem_arquivo) -> numero da linha (numero_linha_arquivo). Quando duas movimentacoes do mesmo "
         "processo tem exatamente a mesma data_dt, a ordem entre elas e definida por arquivo/linha e o processo "
         "e sinalizado com a anomalia ORDEM_CRONOLOGICA_AMBIGUA (movimentacoes RESTITUIDO nao entram nessa "
         "checagem de ambiguidade)."),
    ]


def _colunas_episodios():
    return [
        ("processo", "Texto", "Numero do processo ao qual o episodio pertence.", "Grafia original (nao normalizada)."),
        ("processo_norm", "Texto", "Numero do processo normalizado.", "Chave de casamento com a aba Processos."),
        ("numero_episodio", "Numero inteiro", "Ordem do episodio dentro do processo (1, 2, 3...), na ordem "
         "cronologica em que ocorreram.", "Um processo pode ter varios episodios; apenas o primeiro VALIDO (ou "
         "o primeiro de qualquer tipo, se nao houver valido) define os campos oficiais na aba Processos."),
        ("unidade_origem_norm", "Texto", "Unidade imediatamente anterior a entrada na triagem, normalizada.", "-"),
        ("unidade_origem", "Texto", "Mesma unidade acima, com a grafia oficial do arquivo de classificacao.", "-"),
        ("classificacao_origem", "Texto", "Classificacao (CEDENTE/NEUTRA/DESTINATARIA) da unidade de origem.",
         "Precisa ser CEDENTE para o episodio ser VALIDA."),
        ("data_entrada_triagem", "Data/hora", "Data da movimentacao que levou o processo a unidade de triagem.", "-"),
        ("unidade_triagem_norm", "Texto", "Nome normalizado da unidade de triagem envolvida no episodio.", "-"),
        ("unidade_triagem", "Texto", "Mesma unidade acima, com a grafia original do CSV.", "-"),
        ("indicador_entrada_triagem", "Texto", "Indicador normalizado da movimentacao de entrada na triagem.",
         "Precisa ser REDISTRIBUICAO para o episodio ser VALIDA."),
        ("unidade_destino_norm", "Texto", "Primeira unidade permanente apos a triagem, normalizada.",
         "Vazio quando o episodio e do tipo INCOMPLETA."),
        ("unidade_destino", "Texto", "Mesma unidade acima, com a grafia oficial do arquivo de classificacao.", "-"),
        ("classificacao_destino", "Texto", "Classificacao da unidade de destino.",
         "Precisa ser DESTINATARIA para o episodio ser VALIDA."),
        ("data_destino_equalizacao", "Data/hora", "Data da movimentacao que levou o processo a unidade de "
         "destino.", "Vazio quando o episodio e INCOMPLETA."),
        ("indicador_entrada_destino", "Texto", "Indicador normalizado da movimentacao de entrada no destino.",
         "Precisa ser REDISTRIBUICAO para o episodio ser VALIDA."),
        ("tipo_episodio", "Texto", "Classificacao final do episodio.",
         "VALIDA: origem CEDENTE + destino DESTINATARIA + ambas as entradas por REDISTRIBUICAO. FORA_PADRAO: "
         "ha destino, mas alguma das condicoes acima nao e atendida. INCOMPLETA: nao ha unidade permanente de "
         "destino apos a triagem."),
        ("motivo", "Texto", "Explicacao textual de cada condicao nao atendida, separadas por ';'.",
         "Vazio quando tipo_episodio e VALIDA."),
    ]


def _colunas_comparativo():
    return [
        ("unidade_norm", "Texto", "Nome normalizado da unidade permanente.", "-"),
        ("unidade_original", "Texto", "Nome oficial da unidade, conforme o arquivo de classificacao.", "-"),
        ("classificacao", "Texto", "Classificacao da unidade.", "CEDENTE, NEUTRA ou DESTINATARIA."),
        ("quantidade_magistrados", "Numero decimal", "Quantidade de magistrados da unidade, conforme a coluna "
         "'Quantidade de Magistrados' do arquivo de classificacao.", "Vazio (N/A) quando o arquivo de "
         "classificacao nao possui essa coluna, ou quando o valor informado para a unidade e ausente/invalido "
         "(<= 0 ou nao numerico) -- nesses casos a unidade fica de fora do grupo de medidas ponderadas por "
         "magistrado, mas continua normalmente nas demais colunas desta aba."),
        ("quantidade", "Numero inteiro", "Quantidade de processos distintos contados nesta unidade, no cenario "
         "da linha (ver coluna cenario).", "Inclui unidades com quantidade = 0 (que nao receberam nenhum "
         "processo), para que as medias e medidas de dispersao considerem o universo completo de unidades "
         "permanentes."),
        ("percentual_total", "Numero decimal (%)", "Participacao percentual da unidade no total do cenario.",
         "quantidade / soma de quantidade de todas as unidades do mesmo cenario x 100."),
        ("diferenca_absoluta_media", "Numero decimal", "Diferenca entre a quantidade da unidade e a media geral "
         "do cenario.", "quantidade - media das quantidades de todas as unidades do mesmo cenario (media "
         "simples, cada unidade com o mesmo peso)."),
        ("diferenca_percentual_media", "Numero decimal (%)", "A diferenca absoluta acima, expressa em percentual "
         "da media.", "diferenca_absoluta_media / media x 100."),
        ("processos_por_magistrado", "Numero decimal", "Quantidade de processos da unidade dividida pela sua "
         "quantidade de magistrados -- a carga de trabalho por magistrado da unidade, no cenario da linha.",
         "= quantidade / quantidade_magistrados. Vazio (N/A) quando quantidade_magistrados e vazio. Este e o "
         "grupo de medidas ponderadas pela quantidade de magistrados; ver a secao de regras logo apos esta "
         "tabela."),
        ("diferenca_absoluta_media_magistrado", "Numero decimal", "Diferenca entre processos_por_magistrado da "
         "unidade e a media ponderada do sistema no cenario da linha.",
         "= processos_por_magistrado - media ponderada (total de processos / total de magistrados das unidades "
         "com quantidade de magistrados valida no cenario). Diferente de diferenca_absoluta_media: aqui o "
         "centro de referencia e a media PONDERADA por magistrado, nao a media simples por unidade."),
        ("diferenca_percentual_media_magistrado", "Numero decimal (%)", "A diferenca absoluta acima (por "
         "magistrado), expressa em percentual da media ponderada.",
         "diferenca_absoluta_media_magistrado / media ponderada x 100."),
        ("cenario", "Texto", "Identifica a qual cenario a linha pertence.",
         "COM_EQUALIZACAO (baseado em unidade_final_real) ou SEM_EQUALIZACAO (baseado em "
         "unidade_simulada_sem_equalizacao). Cada unidade aparece duas vezes na aba, uma para cada cenario -- "
         "ver a secao de regras dos cenarios comparativos."),
    ]


def _colunas_fluxos():
    return [
        ("unidade_cedente", "Texto", "Nome oficial da unidade cedente da rota.", "-"),
        ("unidade_destinataria", "Texto", "Nome oficial da unidade destinataria da rota.", "-"),
        ("quantidade", "Numero inteiro", "Quantidade de processos distintos que percorreram esta rota "
         "especifica (mesmo par cedente/destinataria).", "Considera apenas processos com equalizacao_valida = "
         "verdadeiro, usando a unidade cedente/destinataria do episodio principal de cada processo."),
        ("participacao_pct", "Numero decimal (%)", "Participacao percentual da rota no total de processos "
         "equalizados validamente.", "quantidade / soma de quantidade de todas as rotas x 100."),
    ]


def _colunas_anomalias():
    return [
        ("processo", "Texto", "Numero do processo ao qual a anomalia se refere.",
         "Pode ser vazio para a anomalia PROCESSO_SEM_NUMERO."),
        ("tipo_anomalia", "Texto", "Codigo do tipo de anomalia.", "Ver catalogo completo logo abaixo desta tabela."),
        ("descricao", "Texto", "Texto explicativo especifico da ocorrencia.",
         "Quando nao ha detalhe adicional, repete a descricao padrao do tipo de anomalia."),
        ("arquivo_origem", "Texto", "Arquivo CSV de referencia para localizar o registro que originou a "
         "anomalia.", "Para anomalias de nivel de processo (nao de uma linha especifica), referencia a "
         "primeira movimentacao do processo."),
        ("numero_linha_arquivo", "Numero inteiro", "Linha fisica, no arquivo acima, de referencia para a "
         "anomalia.", "Mesma logica de referencia da coluna anterior."),
    ]


CATALOGO_ANOMALIAS = [
    ("PROCESSO_SEM_NUMERO", "Processo sem numero", "O campo #Processo esta vazio na movimentacao."),
    ("PROCESSO_FORMATO_INVALIDO", "Processo com formato aparentemente invalido",
     "O numero do processo nao contem nenhum digito, ou tem menos de 3 caracteres."),
    ("DATA_INVALIDA", "Data invalida", "A coluna Data nao pode ser interpretada como uma data valida."),
    ("UNIDADE_VAZIA", "Unidade vazia", "O campo Orgao Julgador esta vazio na movimentacao."),
    ("UNIDADE_NAO_CLASSIFICADA", "Unidade nao classificada",
     "O orgao julgador da movimentacao nao e uma unidade de triagem nem consta no arquivo de classificacao de "
     "unidades permanentes. A unidade fica listada na aba Unidades_Nao_Classificadas."),
    ("PRIMEIRO_MOVIMENTO_NAO_DISTRIBUICAO", "Primeiro movimento diferente de distribuicao",
     "A primeira movimentacao cronologica do processo (analisada) nao tem indicador DISTRIBUICAO."),
    ("MULTIPLA_DISTRIBUICAO_INICIAL", "Mais de uma distribuicao inicial",
     "O processo tem mais de uma movimentacao com indicador DISTRIBUICAO."),
    ("REDISTRIBUICAO_ANTERIOR_A_DISTRIBUICAO", "Redistribuicao com data anterior a distribuicao",
     "Existe uma movimentacao REDISTRIBUICAO com data anterior a data_caso_novo do processo."),
    ("ORDEM_CRONOLOGICA_AMBIGUA", "Ordem cronologica ambigua entre movimentos na mesma data",
     "Duas ou mais movimentacoes (analisadas) do mesmo processo tem exatamente a mesma data/hora; a ordem entre "
     "elas foi definida por arquivo e numero de linha, mas nao ha garantia de que essa seja a ordem real."),
    ("PASSAGEM_TRIAGEM_INCOMPLETA", "Passagem incompleta pela triagem",
     "Um episodio do processo entrou em uma unidade de triagem, mas nao ha unidade permanente de destino "
     "definitiva depois dela."),
    ("ORIGEM_EQUALIZACAO_NAO_CEDENTE", "Origem da equalizacao nao classificada como cedente",
     "Em um episodio FORA_PADRAO, a unidade anterior a triagem nao esta classificada como CEDENTE."),
    ("DESTINO_EQUALIZACAO_NAO_DESTINATARIA", "Destino da equalizacao nao classificado como destinataria",
     "Em um episodio FORA_PADRAO, a primeira unidade permanente apos a triagem nao esta classificada como "
     "DESTINATARIA."),
    ("ULTIMO_MOVIMENTO_EM_TRIAGEM", "Ultimo movimento do processo ainda em unidade de triagem",
     "A ultima movimentacao (analisada) do processo esta em uma unidade de triagem."),
    ("MULTIPLOS_EPISODIOS_EQUALIZACAO", "Processo com mais de um episodio de passagem pela triagem",
     "O processo tem quantidade_episodios_triagem maior que 1."),
    ("MOVIMENTACAO_DUPLICADA_REMOVIDA", "Movimentacao duplicada removida na consolidacao",
     "A mesma movimentacao (processo, indicador, orgao julgador, data, classe judicial e municipio sede, apos "
     "normalizacao) apareceu mais de uma vez nos arquivos de entrada; apenas a primeira ocorrencia (por "
     "arquivo/linha) foi mantida."),
    ("CLASSE_JUDICIAL_DIVERGENTE", "Classe judicial divergente entre movimentos do mesmo processo",
     "As movimentacoes do processo nao tem todas a mesma classe judicial normalizada."),
]


def _colunas_nao_classificadas():
    return [
        ("unidade_norm", "Texto", "Nome normalizado da unidade nao classificada.", "-"),
        ("unidade", "Texto", "Grafia original da unidade, como apareceu no CSV (primeira ocorrencia).", "-"),
        ("quantidade_movimentos", "Numero inteiro", "Quantidade de movimentacoes (analisadas) registradas nesta "
         "unidade nao classificada.", "Nao inclui movimentacoes RESTITUIDO."),
        ("processos_distintos", "Numero inteiro", "Quantidade de processos distintos que tiveram alguma "
         "movimentacao nesta unidade nao classificada.", "-"),
        ("arquivo_exemplo", "Texto", "Um arquivo CSV de exemplo onde a unidade aparece.",
         "Primeira ocorrencia encontrada; util para localizar rapidamente o registro na origem."),
        ("linha_exemplo", "Numero inteiro", "Linha fisica de exemplo, no arquivo acima, onde a unidade "
         "aparece.", "Mesma logica de referencia da coluna anterior."),
    ]


def gerar_documentacao(caminho_saida: Path = CAMINHO_SAIDA) -> Path:
    data_geracao = datetime.now()
    document = Document()
    _config_pagina_a4(document)

    # ---- Capa ---------------------------------------------------------
    document.add_paragraph()
    document.add_paragraph()
    capa = document.add_paragraph()
    capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa.add_run("Documentacao Tecnica")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = COR_TITULO
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Planilha Consolidada do Sistema de Equalizacao TRT12\n(consolidado_equalizacao.xlsx)")
    run2.font.size = Pt(16)
    subtitulo = document.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run("Descricao de cada aba, coluna e regra de negocio envolvida").italic = True
    document.add_paragraph()
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Documento gerado em: {data_geracao.strftime('%d/%m/%Y %H:%M')}\n")
    info.add_run("Gerado por: scripts/gerar_documentacao_planilhas.py")
    document.add_page_break()

    # ---- 1. Visao geral -------------------------------------------------
    _titulo(document, "1. Visao geral")
    _paragrafo(document, TEXTO_INTRODUCAO)
    document.add_paragraph()
    tabela_abas = document.add_table(rows=1, cols=2)
    tabela_abas.style = "Light Grid Accent 1"
    tabela_abas.rows[0].cells[0].text = "Aba"
    tabela_abas.rows[0].cells[1].text = "Conteudo em uma linha"
    for p in tabela_abas.rows[0].cells[0].paragraphs[0].runs:
        p.bold = True
    _marcar_linha_cabecalho_repetido(tabela_abas.rows[0])
    resumo_abas = [
        ("Resumo", "Indicadores de qualidade e metadados da execucao que gerou o arquivo."),
        ("Processos", "Um processo por linha, com todos os campos derivados da analise de equalizacao."),
        ("Movimentacoes", "Todas as movimentacoes consolidadas, deduplicadas e ordenadas cronologicamente."),
        ("Episodios_Equalizacao", "Auditoria de todos os episodios de passagem por unidade de triagem."),
        ("Comparativo_Unidades", "Quantidade de processos por unidade, nos cenarios com e sem equalizacao."),
        ("Fluxos_Equalizacao", "Rotas cedente -> destinataria dos processos equalizados validamente."),
        ("Anomalias", "Catalogo de anomalias de qualidade de dados encontradas na consolidacao."),
        ("Unidades_Nao_Classificadas", "Unidades citadas nos CSVs sem correspondencia no arquivo de classificacao."),
    ]
    for aba, desc in resumo_abas:
        cels = tabela_abas.add_row().cells
        cels[0].text = aba
        cels[1].text = desc
    document.add_page_break()

    # ---- 2. Convencoes gerais --------------------------------------------
    _titulo(document, "2. Convencoes gerais")
    _paragrafo(document, TEXTO_CONVENCOES)
    document.add_page_break()

    # ---- 3 a 10: abas -----------------------------------------------------
    _aba(document, "3", "Resumo", "-- (exclusiva do Excel consolidado)",
         "Painel de indicadores de qualidade e metadados sobre a execucao que gerou o arquivo -- quantos "
         "registros foram lidos, quantos arquivos, quantas duplicidades e anomalias, e assim por diante.",
         "Uma linha por indicador (formato chave/valor).", _colunas_resumo())

    _aba(document, "4", "Processos", "processos_consolidados.csv / .parquet",
         "Tabela analitica principal: um registro por processo, reunindo o resultado completo da reconstrucao "
         "do historico e da deteccao de equalizacao. E a aba usada para responder 'este processo foi "
         "equalizado, e por que (nao)?'.",
         "Um processo distinto por linha.", _colunas_processos())

    _aba(document, "5", "Movimentacoes", "movimentacoes_consolidadas.csv / .parquet",
         "Todas as movimentacoes dos CSVs de entrada, apos deduplicacao exata e ordenacao cronologica por "
         "processo. E o nivel de detalhe mais granular do sistema -- toda anomalia e todo campo derivado da "
         "aba Processos pode ser rastreado ate uma ou mais linhas desta aba.",
         "Uma movimentacao (linha de CSV) por linha, incluindo as classificadas como RESTITUIDO.",
         _colunas_movimentacoes())

    _aba(document, "6", "Episodios_Equalizacao", "episodios_equalizacao.csv",
         "Auditoria de cada episodio de passagem por unidade de triagem identificado no historico dos "
         "processos -- inclusive os episodios invalidos ou incompletos, e os casos em que um mesmo processo "
         "passou pela triagem mais de uma vez.",
         "Um episodio de triagem por linha (um processo pode ter zero, um ou varios episodios).",
         _colunas_episodios())

    _titulo(document, "6.1 Regra de negocio: deteccao de equalizacao", nivel=3)
    _paragrafo(document, TEXTO_REGRA_EQUALIZACAO)
    document.add_page_break()

    _aba(document, "7", "Comparativo_Unidades", "comparativo_unidades.csv",
         "Quantidade de processos por unidade permanente, lado a lado nos dois cenarios (com e sem "
         "equalizacao), com as medidas de participacao percentual e distancia da media usadas nos graficos "
         "comparativos do dashboard e do relatorio gerencial.",
         "Uma unidade permanente por linha, repetida uma vez para cada cenario (2 linhas por unidade).",
         _colunas_comparativo())

    _titulo(document, "7.1 Regra de negocio: cenarios com e sem equalizacao", nivel=3)
    _paragrafo(document, TEXTO_REGRA_CENARIOS)

    _titulo(document, "7.2 Regra de negocio: medidas ponderadas por magistrado", nivel=3)
    _paragrafo(document, TEXTO_REGRA_MAGISTRADO)
    document.add_page_break()

    _aba(document, "8", "Fluxos_Equalizacao", "fluxos_equalizacao.csv",
         "Quantidade de processos equalizados validamente em cada rota especifica entre uma unidade cedente e "
         "uma unidade destinataria -- a base do diagrama Sankey e da matriz de calor do dashboard.",
         "Um par cedente/destinataria por linha (apenas pares com pelo menos um processo).", _colunas_fluxos())

    _titulo(document, "9. Aba `Anomalias`", nivel=2)
    _paragrafo(document, "Catalogo de todas as anomalias de qualidade de dados encontradas durante a "
               "consolidacao -- desde problemas simples de preenchimento (data invalida, processo sem numero) "
               "ate situacoes especificas da logica de equalizacao (passagem incompleta pela triagem, origem "
               "nao classificada como cedente etc.). Um mesmo processo pode ter varias linhas nesta aba, uma "
               "para cada anomalia encontrada.")
    p = document.add_paragraph()
    p.add_run("Granularidade: ").bold = True
    p.add_run("Uma ocorrencia de anomalia por linha (um processo pode ter zero, uma ou varias).")
    p2 = document.add_paragraph()
    p2.add_run("Tambem disponivel em: ").bold = True
    p2.add_run("anomalias.csv")
    _tabela_colunas(document, _colunas_anomalias())

    _titulo(document, "9.1 Catalogo completo dos tipos de anomalia", nivel=3)
    tabela_cat = document.add_table(rows=1, cols=3)
    tabela_cat.style = "Light Grid Accent 1"
    cabecalhos_cat = ["Codigo (tipo_anomalia)", "Descricao padrao", "Quando ocorre"]
    for i, texto in enumerate(cabecalhos_cat):
        tabela_cat.rows[0].cells[i].text = texto
        for r in tabela_cat.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    _marcar_linha_cabecalho_repetido(tabela_cat.rows[0])
    for codigo, descricao_padrao, quando in CATALOGO_ANOMALIAS:
        cels = tabela_cat.add_row().cells
        cels[0].text = codigo
        for r in cels[0].paragraphs[0].runs:
            r.font.name = "Consolas"
            r.font.color.rgb = COR_CODIGO
        cels[1].text = descricao_padrao
        cels[2].text = quando
    document.add_page_break()

    _aba(document, "10", "Unidades_Nao_Classificadas", "unidades_nao_classificadas.csv",
         "Relacao de unidades (orgaos julgadores) citadas nas movimentacoes dos CSVs de entrada que nao sao "
         "unidades de triagem e nao constam no arquivo de classificacao de unidades permanentes -- ou seja, "
         "unidades que o sistema nao conseguiu classificar como CEDENTE, NEUTRA ou DESTINATARIA. O sistema "
         "nunca atribui uma classificacao automaticamente por aproximacao; esta aba existe para apoiar a "
         "revisao manual do arquivo de classificacao.",
         "Uma unidade nao classificada por linha (agregando todas as suas ocorrencias).",
         _colunas_nao_classificadas())

    _rodape(document, data_geracao)
    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(caminho_saida))
    return caminho_saida


if __name__ == "__main__":
    caminho = gerar_documentacao()
    print(f"Documentacao gerada em: {caminho}")
