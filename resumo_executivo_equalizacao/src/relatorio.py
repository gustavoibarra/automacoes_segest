"""Monta o DOCX do Resumo Executivo do Sistema de Equalizacao, seguindo a
estrutura e a identidade visual do documento de referencia (ver
config.py): capa com faixa de KPIs, secoes 1-5 em retrato (Carta/Letter) e
3 anexos em paisagem, com cabecalho (logo) e rodape (paginacao) em todas as
paginas.
"""
from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config import (
    COR_BORDA_TABELA,
    COR_KPI_REDUCAO,
    COR_KPI_VOLUME,
    COR_NOTA_FUNDO,
    COR_PERIODO,
    COR_RODAPE_TEXTO,
    COR_SUBTITULO,
    COR_TABELA_CABECALHO,
    COR_TITULO,
    FONTE_PADRAO,
    LOGO_PATH,
    PAGINA_ALTURA_CM,
    PAGINA_LARGURA_CM,
    SUBTITULO_RELATORIO,
    TITULO_RELATORIO,
)
from src import dados, graficos, metricas
from src.formatting import (
    formatar_data_extenso_dia1,
    formatar_numero,
    formatar_percentual,
    formatar_percentual_com_sinal,
)

_RGB_TITULO = RGBColor.from_string(COR_TITULO)
_RGB_SUBTITULO = RGBColor.from_string(COR_SUBTITULO)
_RGB_PERIODO = RGBColor.from_string(COR_PERIODO)
_RGB_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
_RGB_RODAPE = RGBColor.from_string(COR_RODAPE_TEXTO)


# ---------------------------------------------------------------------------
# Utilitarios de baixo nivel (pagina, cabecalho, rodape, sombreamento)
# ---------------------------------------------------------------------------

def _definir_tamanho_secao(section, largura_cm: float, altura_cm: float, margens_lr: float, margens_tb: float) -> None:
    section.page_width = Cm(largura_cm)
    section.page_height = Cm(altura_cm)
    section.left_margin = Cm(margens_lr)
    section.right_margin = Cm(margens_lr)
    section.top_margin = Cm(margens_tb)
    section.bottom_margin = Cm(margens_tb)


def _sombrear_celula(cell, cor_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


def _definir_bordas_tabela(tabela, cor_hex: str = COR_BORDA_TABELA, tamanho: int = 4) -> None:
    """Aplica bordas finas e continuas (externas e internas) na cor informada."""
    tblPr = tabela._tbl.tblPr
    bordas = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borda = OxmlElement(f"w:{tag}")
        borda.set(qn("w:val"), "single")
        borda.set(qn("w:sz"), str(tamanho))
        borda.set(qn("w:space"), "0")
        borda.set(qn("w:color"), cor_hex)
        bordas.append(borda)
    tblPr.append(bordas)


def _marcar_cabecalho_repetido(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _definir_cabecalho_logo(document: Document, section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragrafo = header.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        run = paragrafo.add_run()
        run.add_picture(str(LOGO_PATH), height=Cm(1.1))


def _adicionar_campo(paragraph, tipo: str) -> None:
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {tipo} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


def _definir_rodape(document: Document, section, data_geracao: datetime) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragrafo = footer.paragraphs[0]
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.text = ""
    run = paragrafo.add_run("Relatório Executivo do Sistema de Equalização | Página ")
    run.font.name, run.font.size, run.font.color.rgb = FONTE_PADRAO, Pt(10.5), _RGB_RODAPE
    _adicionar_campo(paragrafo, "PAGE")
    run2 = paragrafo.add_run(" de ")
    run2.font.name, run2.font.size, run2.font.color.rgb = FONTE_PADRAO, Pt(10.5), _RGB_RODAPE
    _adicionar_campo(paragrafo, "NUMPAGES")
    run3 = paragrafo.add_run(f" | Gerado em {data_geracao.strftime('%d/%m/%Y %H:%M')}")
    run3.font.name, run3.font.size, run3.font.color.rgb = FONTE_PADRAO, Pt(10.5), _RGB_RODAPE


def _nova_secao_paisagem(document: Document, data_geracao: datetime):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    _definir_tamanho_secao(section, PAGINA_ALTURA_CM, PAGINA_LARGURA_CM, margens_lr=1.07, margens_tb=1.14)
    _definir_cabecalho_logo(document, section)
    _definir_rodape(document, section, data_geracao)
    return section


# ---------------------------------------------------------------------------
# Blocos de texto
# ---------------------------------------------------------------------------

def _titulo_secao(document: Document, texto: str, nivel_esboco: int = 0):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(nivel_esboco))
    pPr.append(outline)
    run = p.add_run(texto)
    run.bold = True
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(18)
    run.font.color.rgb = _RGB_TITULO
    return p


def _subtitulo_secao(document: Document, texto: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.bold = True
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(13)
    run.font.color.rgb = _RGB_TITULO
    return p


def _paragrafo(document: Document, texto: str, tamanho: float = 11, cor: RGBColor | None = None,
               negrito: bool = False):
    p = document.add_paragraph()
    run = p.add_run(texto)
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(tamanho)
    run.bold = negrito
    if cor is not None:
        run.font.color.rgb = cor
    return p


def _caixa_nota(document: Document, texto: str, prefixo: str = "Nota: "):
    tabela = document.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    _definir_bordas_tabela(tabela)
    celula = tabela.rows[0].cells[0]
    celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _sombrear_celula(celula, COR_NOTA_FUNDO)
    p = celula.paragraphs[0]
    if prefixo:
        run_prefixo = p.add_run(prefixo)
        run_prefixo.bold = True
        run_prefixo.font.name = FONTE_PADRAO
        run_prefixo.font.size = Pt(9)
        run_prefixo.font.color.rgb = _RGB_TITULO
    run = p.add_run(texto)
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(9)
    run.font.color.rgb = _RGB_TITULO
    document.add_paragraph()
    return tabela


def _kpi_band(document: Document, cards: list[tuple[str, str, str]]) -> None:
    """cards: lista de (valor_formatado, rotulo, cor_hex)."""
    tabela = document.add_table(rows=1, cols=len(cards))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    _definir_bordas_tabela(tabela)
    largura = Cm((PAGINA_LARGURA_CM - 2 * 1.65) / len(cards))
    for i, (valor, rotulo, cor) in enumerate(cards):
        celula = tabela.rows[0].cells[i]
        celula.width = largura
        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _sombrear_celula(celula, cor)
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_valor = p.add_run(valor + "\n")
        run_valor.bold = True
        run_valor.font.name = FONTE_PADRAO
        run_valor.font.size = Pt(16 if len(cards) <= 3 or i < 3 else 12)
        run_valor.font.color.rgb = _RGB_BRANCO
        run_rotulo = p.add_run(rotulo)
        run_rotulo.font.name = FONTE_PADRAO
        run_rotulo.font.size = Pt(8 if i < 3 else 7)
        run_rotulo.font.color.rgb = _RGB_BRANCO
    document.add_paragraph()


def _celula_texto(celula, texto: str, negrito: bool = False, cor: RGBColor | None = None,
                   tamanho: float = 9.5, alinhamento=None) -> None:
    celula.text = ""
    celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = celula.paragraphs[0]
    if alinhamento is not None:
        p.alignment = alinhamento
    run = p.add_run(str(texto))
    run.bold = negrito
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(tamanho)
    if cor is not None:
        run.font.color.rgb = cor


def _tabela_simples(document: Document, cabecalho: list[str], linhas: list[list[str]],
                     larguras_cm: list[float] | None = None, negrito_linhas: set[int] | None = None,
                     colunas_quantidade: set[int] | None = None) -> None:
    """colunas_quantidade: indices de colunas numericas, centralizadas horizontalmente."""
    negrito_linhas = negrito_linhas or set()
    colunas_quantidade = colunas_quantidade or set()
    tabela = document.add_table(rows=1, cols=len(cabecalho))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = larguras_cm is None
    _definir_bordas_tabela(tabela)
    for i, texto in enumerate(cabecalho):
        celula = tabela.rows[0].cells[i]
        _sombrear_celula(celula, COR_TABELA_CABECALHO)
        alinhamento = WD_ALIGN_PARAGRAPH.CENTER if i in colunas_quantidade else None
        _celula_texto(celula, texto, negrito=True, cor=_RGB_BRANCO, alinhamento=alinhamento)
        if larguras_cm:
            celula.width = Cm(larguras_cm[i])
    _marcar_cabecalho_repetido(tabela.rows[0])

    for indice, linha in enumerate(linhas):
        celulas = tabela.add_row().cells
        eh_negrito = indice in negrito_linhas
        for i, valor in enumerate(linha):
            alinhamento = WD_ALIGN_PARAGRAPH.CENTER if i in colunas_quantidade else None
            _celula_texto(celulas[i], valor, negrito=eh_negrito, alinhamento=alinhamento)
            if larguras_cm:
                celulas[i].width = Cm(larguras_cm[i])
    document.add_paragraph()


def _tabela_cabecalho_agrupado(document: Document, grupos: list[tuple[str, int]], subcabecalho: list[str],
                                linhas: list[list[str]], larguras_cm: list[float],
                                negrito_linhas: set[int] | None = None,
                                colunas_quantidade: set[int] | None = None) -> None:
    """Tabela com 2 linhas de cabecalho: a primeira com celulas mescladas
    (grupos de colunas), a segunda com os rotulos individuais. Quando uma
    celula de grupo (span=1) repete o rotulo do subcabecalho correspondente,
    a celula de cima fica vazia para nao duplicar o texto. Por padrao, todas
    as colunas exceto a primeira (rotulo) sao tratadas como quantidades e
    centralizadas horizontalmente."""
    negrito_linhas = negrito_linhas or set()
    colunas_quantidade = (colunas_quantidade if colunas_quantidade is not None
                           else set(range(1, len(subcabecalho))))
    n_cols = len(subcabecalho)
    tabela = document.add_table(rows=2, cols=n_cols)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    _definir_bordas_tabela(tabela)

    col = 0
    for texto, span in grupos:
        if span > 1:
            celula = tabela.cell(0, col).merge(tabela.cell(0, col + span - 1))
        else:
            celula = tabela.cell(0, col)
        _sombrear_celula(celula, COR_TABELA_CABECALHO)
        texto_exibido = "" if span == 1 and texto == subcabecalho[col] else texto
        _celula_texto(celula, texto_exibido, negrito=True, cor=_RGB_BRANCO, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
        col += span

    for i, texto in enumerate(subcabecalho):
        celula = tabela.cell(1, i)
        _sombrear_celula(celula, COR_TABELA_CABECALHO)
        alinhamento = WD_ALIGN_PARAGRAPH.CENTER if i in colunas_quantidade else None
        _celula_texto(celula, texto, negrito=True, cor=_RGB_BRANCO, tamanho=9, alinhamento=alinhamento)
        celula.width = Cm(larguras_cm[i])

    _marcar_cabecalho_repetido(tabela.rows[0])
    _marcar_cabecalho_repetido(tabela.rows[1])

    for indice, linha in enumerate(linhas):
        celulas = tabela.add_row().cells
        eh_negrito = indice in negrito_linhas
        for i, valor in enumerate(linha):
            alinhamento = WD_ALIGN_PARAGRAPH.CENTER if i in colunas_quantidade else None
            _celula_texto(celulas[i], valor, negrito=eh_negrito, tamanho=9, alinhamento=alinhamento)
            celulas[i].width = Cm(larguras_cm[i])
    document.add_paragraph()


def _adicionar_figura(document: Document, fig, tmp_dir: Path, nome_arquivo: str, legenda: str,
                       largura_cm: float = 16.0) -> None:
    caminho = tmp_dir / nome_arquivo
    fig.savefig(caminho, dpi=150, bbox_inches="tight")
    import matplotlib.pyplot as plt
    plt.close(fig)
    document.add_picture(str(caminho), width=Cm(largura_cm))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legenda_p = document.add_paragraph()
    run = legenda_p.add_run(legenda)
    run.italic = True
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(9.5)
    run.font.color.rgb = _RGB_SUBTITULO


# ---------------------------------------------------------------------------
# Orquestracao
# ---------------------------------------------------------------------------

def gerar_relatorio(caminho_arquivo: Path, caminho_saida: Path, *, inicio: pd.Timestamp | None = None,
                     fim: pd.Timestamp | None = None, data_geracao: datetime | None = None) -> Path:
    data_geracao = data_geracao or datetime.now()

    base = dados.carregar_consolidado(caminho_arquivo)
    unidades_df = dados.unidades_mestre(base["comparativo"])
    inicio, fim = dados.resolver_periodo(base["processos"], inicio, fim)

    casos_novos_df = dados.casos_novos_periodo(base["processos"], inicio, fim)
    movimentados_df = dados.processos_movimentados_periodo(base["processos"], base["movimentacoes"], inicio, fim)

    tabela_com = metricas.cenario_com_equalizacao(unidades_df, movimentados_df)
    tabela_sem = metricas.cenario_sem_equalizacao(unidades_df, movimentados_df)

    kpis = metricas.kpis(casos_novos_df, tabela_com, tabela_sem)
    dim_unidade = metricas.dimensao_dispersao(tabela_com, tabela_sem, "quantidade")

    media_pond_com = metricas.media_ponderada_por_magistrado(tabela_com)
    media_pond_sem = metricas.media_ponderada_por_magistrado(tabela_sem)
    dim_magistrado = metricas.dimensao_dispersao(
        tabela_com, tabela_sem, "processos_por_magistrado",
        media_referencia_com=media_pond_com, media_referencia_sem=media_pond_sem,
    )

    grupos_df = metricas.tabela_grupos(unidades_df, tabela_com, tabela_sem)
    convergencia = metricas.convergencia_cedentes_destinatarias(grupos_df)

    anexo_cedentes = metricas.tabela_anexo_cedentes(unidades_df, casos_novos_df, tabela_com, tabela_sem)
    anexo_destinatarias = metricas.tabela_anexo_destinatarias(unidades_df, casos_novos_df, tabela_com, tabela_sem)

    document = Document()
    section0 = document.sections[0]
    section0.orientation = WD_ORIENT.PORTRAIT
    _definir_tamanho_secao(section0, PAGINA_LARGURA_CM, PAGINA_ALTURA_CM, margens_lr=1.65, margens_tb=1.40)
    _definir_cabecalho_logo(document, section0)
    _definir_rodape(document, section0, data_geracao)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        _montar_capa(document, kpis, inicio, fim)
        _montar_secao1_criterios(document)
        _montar_secao2_resultado_geral(document, kpis, tabela_com, unidades_df)
        _montar_secao3_dimensao_unidade(document, dim_unidade, tmp_dir)
        _montar_secao4_dimensao_magistrado(document, dim_magistrado, tmp_dir)
        _montar_secao5_grupos(document, grupos_df, convergencia, tmp_dir)

        _nova_secao_paisagem(document, data_geracao)
        _montar_anexo_unidades(document, "Anexo 1. Unidades cedentes", "cedente", "enviados",
                                anexo_cedentes, grupos_df, metricas.CEDENTE)
        _montar_anexo_unidades(document, "Anexo 2. Unidades destinatárias", "destinatária", "recebidos",
                                anexo_destinatarias, grupos_df, metricas.DESTINATARIA)
        _montar_anexo3_resumo_grupo(document, grupos_df, convergencia)

    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(caminho_saida))
    return caminho_saida


def _montar_capa(document: Document, kpis: dict, inicio: pd.Timestamp, fim: pd.Timestamp) -> None:
    document.add_paragraph()
    titulo = document.add_paragraph()
    run = titulo.add_run(TITULO_RELATORIO)
    run.bold = True
    run.font.name = FONTE_PADRAO
    run.font.size = Pt(25)
    run.font.color.rgb = _RGB_TITULO

    subtitulo = document.add_paragraph()
    run_sub = subtitulo.add_run(SUBTITULO_RELATORIO)
    run_sub.font.name = FONTE_PADRAO
    run_sub.font.size = Pt(11.5)
    run_sub.font.color.rgb = _RGB_SUBTITULO

    periodo = document.add_paragraph()
    run_per = periodo.add_run(
        f"Período analisado: {formatar_data_extenso_dia1(inicio)} a {formatar_data_extenso_dia1(fim)}"
    )
    run_per.bold = True
    run_per.font.name = FONTE_PADRAO
    run_per.font.size = Pt(11)
    run_per.font.color.rgb = _RGB_PERIODO
    document.add_paragraph()

    _kpi_band(document, [
        (formatar_numero(kpis["casos_novos"]), "casos novos", COR_KPI_VOLUME),
        (formatar_numero(kpis["processos_equalizados"]), "processos equalizados", COR_KPI_VOLUME),
        (formatar_percentual(kpis["percentual_equalizados"]), "dos casos novos", COR_KPI_VOLUME),
        (formatar_percentual(kpis["reducao_cv_unidade_pct"]), "redução relativa do CV\npor unidade", COR_KPI_REDUCAO),
        (formatar_percentual(kpis["reducao_cv_magistrado_pct"]), "redução relativa do CV\npor magistrado", COR_KPI_REDUCAO),
    ])

    _paragrafo(
        document,
        "Este relatório apresenta uma leitura executiva dos processos distribuídos e equalizados no período. "
        "A análise é feita em duas dimensões complementares: a quantidade de processos em cada unidade e a "
        "carga de trabalho medida pela quantidade de processos por magistrado."
    )


def _montar_secao1_criterios(document: Document) -> None:
    _titulo_secao(document, "1. Critérios de leitura", nivel_esboco=1)
    _paragrafo(
        document,
        "Um processo é considerado equalizado quando há movimentação de unidade cedente para unidade "
        "destinatária para fins de equalização, passando pelas unidades de triagem e com movimentos de "
        "redistribuição."
    )
    _paragrafo(
        document,
        "A principal medida para análise de eficiência do sistema é o coeficiente de variação, que mede o "
        "tamanho da desigualdade entre as unidades em comparação com o nível médio de processos ou de carga "
        "de trabalho. Para tanto, foi utilizada a seguinte escala:"
    )
    from config import FAIXAS_CV
    linhas_cv = []
    for minimo, maximo, rotulo, _cor in FAIXAS_CV:
        if minimo == 0:
            faixa = f"0% a {formatar_numero(maximo)}%"
        elif maximo == float("inf"):
            faixa = f"Acima de {formatar_numero(minimo)}%"
        else:
            faixa = f"Acima de {formatar_numero(minimo)}% a {formatar_numero(maximo)}%"
        linhas_cv.append([faixa, rotulo])
    _tabela_simples(document, ["CV", "Classificação utilizada"], linhas_cv, larguras_cm=[6, 10],
                     colunas_quantidade={0})

    _paragrafo(
        document,
        "Além disso, as medidas são apresentadas em duas dimensões, comparando, quando possível, o cenário "
        "“com equalização” e “sem equalização”:"
    )
    _tabela_simples(document, ["Dimensão", "Medida principal", "Leitura"], [
        ["Quantidade por unidade", "Número de processos",
         "Compara quanto cada unidade concentra em relação à média das unidades."],
        ["Carga de trabalho", "Processos por magistrado",
         "Divide a quantidade da unidade pelo número de magistrados e compara com a média ponderada do sistema."],
    ], larguras_cm=[4.5, 4.5, 8])

    _caixa_nota(
        document,
        "os cenários “sem equalização” são projeções que representam como seria a situação das "
        "unidades em termos de volume de processos e de carga de trabalho dos magistrados caso o sistema de "
        "equalização não tivesse sido implantado no período."
    )
    document.add_page_break()


def _montar_secao2_resultado_geral(document: Document, kpis: dict, tabela_com: pd.DataFrame,
                                    unidades_df: pd.DataFrame) -> None:
    _titulo_secao(document, "2. Resultado geral do período", nivel_esboco=1)
    _paragrafo(
        document,
        f"Ingressaram no período {formatar_numero(kpis['casos_novos'])} casos novos. Destes, "
        f"{formatar_numero(kpis['processos_equalizados'])} processos foram protocolados em unidades cedentes "
        f"e foram equalizados, o que equivale a {formatar_percentual(kpis['percentual_equalizados'])}."
    )
    # Soma da tabela por unidade (nao a contagem bruta de processos_movimentados_periodo): assim,
    # processos cuja unidade final/simulada nao pertence ao universo de unidades permanentes
    # classificadas (ex.: orgaos de 2o grau) ficam de fora, igual as demais tabelas comparativas.
    total_movimentados = int(tabela_com["quantidade"].sum())
    qtd_unidades = int(len(unidades_df))
    qtd_magistrados = unidades_df["quantidade_magistrados"].sum(skipna=True)
    _paragrafo(
        document,
        f"As análises comparativas consideram {formatar_numero(total_movimentados)} processos movimentados em "
        f"{formatar_numero(qtd_unidades)} unidades e {formatar_numero(qtd_magistrados)} magistrados(as) com "
        "lotação efetiva nas unidades judiciárias de primeiro grau, conforme o Provimento CR nº 03/2026."
    )
    diferenca = total_movimentados - kpis["casos_novos"]
    if diferenca != 0:
        _caixa_nota(
            document,
            f"a diferença entre o número de casos novos ({formatar_numero(kpis['casos_novos'])}) e a "
            f"quantidade de processos movimentados ({formatar_numero(total_movimentados)}) decorre de "
            "situações transitórias, inerentes ao processo de implantação do Sistema de Equalização, em que "
            "ocorreram redistribuições de processos ingressados antes do período analisado, mas que "
            "influenciam os quantitativos utilizados nas análises comparativas."
        )
    document.add_page_break()


def _tabela_medidas_dimensao(document: Document, dim: dict, unidade_medida: str) -> None:
    sem, com = dim["sem"], dim["com"]
    _tabela_simples(document, ["Medida", "Sem equalização", "Com equalização", "Variação"], [
        [f"Média por {unidade_medida}", formatar_numero(sem["media"], 2), formatar_numero(com["media"], 2),
         "sem alteração" if abs(sem["media"] - com["media"]) < 1e-9 else formatar_percentual_com_sinal(
             (com["media"] - sem["media"]) / sem["media"] * 100 if sem["media"] else float("nan"))],
        ["Desvio-padrão", formatar_numero(sem["desvio_padrao"], 2), formatar_numero(com["desvio_padrao"], 2),
         formatar_percentual_com_sinal(-dim["reducao_desvio_padrao_pct"])],
        ["Coeficiente de variação", formatar_percentual(sem["coeficiente_variacao"], 2),
         formatar_percentual(com["coeficiente_variacao"], 2),
         f"{formatar_numero(-dim['reducao_cv_pp'], 2)} p.p. ({formatar_percentual_com_sinal(-dim['reducao_cv_pct'])})"],
        ["Faixa do coeficiente de variação", dim["faixa_sem"], dim["faixa_com"],
         "melhora de faixa" if dim["melhora"] else ("piora de faixa" if dim["faixa_sem"] != dim["faixa_com"] else "sem alteração")],
        ["Menor e maior quantidade", f"{formatar_numero(sem['minimo'])} a {formatar_numero(sem['maximo'])}",
         f"{formatar_numero(com['minimo'])} a {formatar_numero(com['maximo'])}", "intervalo menor" if com["amplitude"] < sem["amplitude"] else "intervalo maior"],
        ["Amplitude", formatar_numero(sem["amplitude"], 2), formatar_numero(com["amplitude"], 2),
         formatar_percentual_com_sinal(-dim["reducao_amplitude_pct"])],
    ], larguras_cm=[5, 4, 4, 5.5], colunas_quantidade={1, 2, 3})


def _montar_secao3_dimensao_unidade(document: Document, dim: dict, tmp_dir: Path) -> None:
    _titulo_secao(document, "3. Dimensão 1 — quantidade de processos por unidade", nivel_esboco=1)
    media = dim["sem"]["media"]
    _paragrafo(
        document,
        f"A média geral do estado no período foi de {formatar_numero(media, 2)} processos por unidade nos "
        "dois cenários, “sem equalização” e “com equalização”. Como o total não muda, por "
        "se tratar do volume estadual de processos, o principal indicador de comparação utilizado é o "
        "coeficiente de variação, que expressa a dispersão relativa das quantidades entre as unidades."
    )
    fig = graficos.grafico_cv_faixas(dim["sem"]["coeficiente_variacao"], dim["com"]["coeficiente_variacao"],
                                      "Coeficiente de variação da quantidade de processos")
    _adicionar_figura(document, fig, tmp_dir, "figura1.png",
                       "Figura 1. Coeficiente de variação da quantidade de processos nos cenários sem e com equalização.")
    _tabela_medidas_dimensao(document, dim, "unidade")
    _caixa_nota(
        document,
        f"O coeficiente de variação caiu de {formatar_percentual(dim['sem']['coeficiente_variacao'], 2)} para "
        f"{formatar_percentual(dim['com']['coeficiente_variacao'], 2)}, redução de "
        f"{formatar_numero(dim['reducao_cv_pp'], 2)} pontos percentuais ou "
        f"{formatar_percentual(dim['reducao_cv_pct'], 1)} em termos relativos. Pela faixa de leitura adotada, "
        f"a distribuição passou de {dim['faixa_sem'].lower()} para {dim['faixa_com'].lower()}. O desvio-padrão "
        f"também caiu {formatar_percentual(dim['reducao_desvio_padrao_pct'], 1)}, e a amplitude entre a menor "
        f"e a maior unidade caiu {formatar_percentual(dim['reducao_amplitude_pct'], 1)}.",
        prefixo="Indicador principal na dimensão quantitativa: coeficiente de variação. ",
    )
    document.add_page_break()


def _montar_secao4_dimensao_magistrado(document: Document, dim: dict, tmp_dir: Path) -> None:
    _titulo_secao(document, "4. Dimensão 2 — carga de trabalho por magistrado", nivel_esboco=1)
    media = dim["sem"]["media"]
    _paragrafo(
        document,
        "A carga de trabalho é calculada dividindo a quantidade de processos de cada unidade pelo número de "
        "magistrados lotados, conforme estabelecido no Provimento CR nº 03/2026. A média ponderada do sistema "
        f"é de {formatar_numero(media, 2)} processos por magistrado. Essa média permanece igual nos dois "
        "cenários, mas o coeficiente de variação mostra quanto as cargas das unidades se afastam "
        "proporcionalmente desse patamar."
    )
    fig = graficos.grafico_cv_faixas(dim["sem"]["coeficiente_variacao"], dim["com"]["coeficiente_variacao"],
                                      "Coeficiente de variação da carga de trabalho por magistrado")
    _adicionar_figura(document, fig, tmp_dir, "figura2.png",
                       "Figura 2. Coeficiente de variação da carga de trabalho por magistrado nos cenários sem e com equalização.")
    _tabela_medidas_dimensao(document, dim, "magistrado")
    _caixa_nota(
        document,
        f"O coeficiente de variação caiu de {formatar_percentual(dim['sem']['coeficiente_variacao'], 2)} para "
        f"{formatar_percentual(dim['com']['coeficiente_variacao'], 2)}, redução de "
        f"{formatar_numero(dim['reducao_cv_pp'], 2)} pontos percentuais ou "
        f"{formatar_percentual(dim['reducao_cv_pct'], 1)} em termos relativos. A classificação passou de "
        f"{dim['faixa_sem'].lower()} para {dim['faixa_com'].lower()}. O desvio-padrão caiu "
        f"{formatar_percentual(dim['reducao_desvio_padrao_pct'], 1)}, e a amplitude da carga caiu de "
        f"{formatar_numero(dim['sem']['amplitude'], 1)} para {formatar_numero(dim['com']['amplitude'], 1)} "
        f"processos por magistrado, redução de {formatar_percentual(dim['reducao_amplitude_pct'], 1)}.",
        prefixo="Indicador principal na dimensão quantitativa: coeficiente de variação. ",
    )
    document.add_page_break()


_ROTULO_GRUPO_TEXTO = {"CEDENTE": "Cedentes", "NEUTRA": "Neutras", "DESTINATARIA": "Destinatárias"}


def _montar_secao5_grupos(document: Document, grupos_df: pd.DataFrame, convergencia: dict, tmp_dir: Path) -> None:
    _titulo_secao(document, "5. Comparação entre cedentes, destinatárias e neutras", nivel_esboco=1)
    _paragrafo(
        document,
        "A comparação por grupo evidencia a mudança entre as unidades que enviaram processos, as que "
        "receberam processos e as unidades neutras. Os gráficos mostram a média de processos por unidade e a "
        "carga média por magistrado."
    )
    fig = graficos.grafico_comparativo_duplo(grupos_df)
    _adicionar_figura(document, fig, tmp_dir, "figura3.png",
                       "Figura 3. Média de processos por unidade e carga média por magistrado, por grupo.", largura_cm=17)

    grupos_ordenados = grupos_df[grupos_df["grupo"] != "Total do sistema"]
    linhas = []
    for _, linha in grupos_ordenados.iterrows():
        linhas.append([
            _ROTULO_GRUPO_TEXTO.get(linha["grupo"], linha["grupo"]),
            formatar_numero(linha["quantidade_unidades"]), formatar_numero(linha["quantidade_magistrados"]),
            formatar_numero(linha["total_sem"]), formatar_numero(linha["media_unidade_sem"], 2), formatar_numero(linha["media_magistrado_sem"], 2),
            formatar_numero(linha["total_com"]), formatar_numero(linha["media_unidade_com"], 2), formatar_numero(linha["media_magistrado_com"], 2),
        ])
    _tabela_cabecalho_agrupado(
        document,
        grupos=[("Grupo", 1), ("Unid.", 1), ("Mag.", 1), ("Sem equalização", 3), ("Com equalização", 3)],
        subcabecalho=["Grupo", "Unid.", "Mag.", "Total", "Proc./unid.", "Proc./mag.", "Total", "Proc./unid.", "Proc./mag."],
        linhas=linhas, larguras_cm=[3, 1.6, 1.6, 1.8, 2, 2, 1.8, 2, 2],
    )

    cedente_row = grupos_df[grupos_df["grupo"] == metricas.CEDENTE].iloc[0]
    destinataria_row = grupos_df[grupos_df["grupo"] == metricas.DESTINATARIA].iloc[0]
    conv_unid = convergencia["unidade"]
    _caixa_nota(
        document,
        f"A diferença entre as médias de cedentes e destinatárias caiu de {formatar_numero(conv_unid['sem'], 2)} "
        f"para {formatar_numero(conv_unid['com'], 2)} processos por unidade, redução de "
        f"{formatar_percentual(conv_unid['reducao_relativa_pct'], 1)}.",
        prefixo="Quantidade por unidade: ",
    )
    conv_mag = convergencia["magistrado"]
    _caixa_nota(
        document,
        f"A diferença entre cedentes e destinatárias caiu de {formatar_numero(conv_mag['sem'], 2)} para "
        f"{formatar_numero(conv_mag['com'], 2)} processos por magistrado, redução de "
        f"{formatar_percentual(conv_mag['reducao_relativa_pct'], 1)}.",
        prefixo="Quantidade por magistrado: ",
    )
    document.add_page_break()


def _montar_anexo_unidades(document: Document, titulo: str, rotulo_singular: str, rotulo_fluxo: str,
                            tabela_anexo: pd.DataFrame, grupos_df: pd.DataFrame, grupo: str) -> None:
    _titulo_secao(document, titulo, nivel_esboco=1)
    _paragrafo(
        document,
        "Fluxo do período e posição dos processos nos cenários sem e com equalização. Casos novos são "
        "atribuídos pela unidade de distribuição inicial."
    )
    linhas = []
    for _, linha in tabela_anexo.iterrows():
        linhas.append([
            linha["unidade"], formatar_numero(linha["quantidade_magistrados"]),
            formatar_numero(linha["casos_novos"]), formatar_numero(linha["equalizados"]),
            formatar_numero(linha["sem_processos"]), formatar_numero(linha["sem_proc_magistrado"], 1),
            formatar_numero(linha["com_processos"]), formatar_numero(linha["com_proc_magistrado"], 1),
            formatar_numero(linha["variacao_processos"]), formatar_numero(linha["variacao_proc_magistrado"], 1),
        ])

    totais = metricas.totais_anexo(tabela_anexo)
    linhas.append([
        "TOTAL", formatar_numero(totais["quantidade_magistrados"]), formatar_numero(totais["casos_novos"]),
        formatar_numero(totais["equalizados"]), formatar_numero(totais["sem_processos"]), "—",
        formatar_numero(totais["com_processos"]), "—", formatar_numero(totais["variacao_processos"]), "—",
    ])
    media = metricas.media_anexo(tabela_anexo, grupos_df, grupo)
    linhas.append([
        "MÉDIA POR UNIDADE / CARGA DO GRUPO", formatar_numero(media["quantidade_magistrados"], 1),
        formatar_numero(media["casos_novos"], 1), formatar_numero(media["equalizados"], 1),
        formatar_numero(media["sem_processos"], 1), formatar_numero(media["sem_proc_magistrado"], 1),
        formatar_numero(media["com_processos"], 1), formatar_numero(media["com_proc_magistrado"], 1),
        formatar_numero(media["variacao_processos"], 1), formatar_numero(media["variacao_proc_magistrado"], 1),
    ])

    _tabela_cabecalho_agrupado(
        document,
        grupos=[("Unidade", 1), ("Mag.", 1), ("Casos novos", 1), (f"Equalizados {rotulo_fluxo}", 1),
                ("Sem equalização", 2), ("Com equalização", 2), ("Variação líquida", 2)],
        subcabecalho=["Unidade", "Mag.", "Casos novos", f"Equaliz. {rotulo_fluxo}", "Processos", "Proc./mag.",
                      "Processos", "Proc./mag.", "Processos", "Proc./mag."],
        linhas=linhas, larguras_cm=[6.5, 1.6, 2, 2.2, 2, 2, 2, 2, 2, 2],
        negrito_linhas={len(linhas) - 2, len(linhas) - 1},
    )
    document.add_page_break()


def _montar_anexo3_resumo_grupo(document: Document, grupos_df: pd.DataFrame, convergencia: dict) -> None:
    _titulo_secao(document, "Anexo 3 – Resumo por Grupo de Unidade", nivel_esboco=1)
    _paragrafo(document, "Comparação dos cenários com e sem equalização")
    _subtitulo_secao(document, "Quadro 1 – Síntese das medidas por grupo")

    linhas = []
    for _, linha in grupos_df.iterrows():
        rotulo = _ROTULO_GRUPO_TEXTO.get(linha["grupo"], linha["grupo"])
        eh_total = linha["grupo"] == "Total do sistema"
        linhas.append([
            rotulo, formatar_numero(linha["quantidade_unidades"]), formatar_numero(linha["quantidade_magistrados"]),
            formatar_numero(linha["total_sem"]), formatar_numero(linha["media_unidade_sem"], 2), formatar_numero(linha["media_magistrado_sem"], 2),
            formatar_numero(linha["total_com"]), formatar_numero(linha["media_unidade_com"], 2), formatar_numero(linha["media_magistrado_com"], 2),
            formatar_numero(linha["variacao_processos"]) if not eh_total else "0",
            (formatar_percentual_com_sinal(linha["variacao_pct"], 1) if pd.notna(linha["variacao_pct"]) else "N/A") if not eh_total else "0,0%",
        ])
    _tabela_cabecalho_agrupado(
        document,
        grupos=[("Grupo", 1), ("Estrutura", 2), ("Cenário sem equalização", 3), ("Cenário com equalização", 3), ("Variação líquida", 2)],
        subcabecalho=["Grupo", "Unidades", "Magistrados", "Total", "Média/unid.", "Média/mag.",
                      "Total", "Média/unid.", "Média/mag.", "Processos", "%"],
        linhas=linhas, larguras_cm=[3, 2, 2.2, 1.8, 2, 2, 1.8, 2, 2, 2, 2],
        negrito_linhas={len(linhas) - 1},
    )

    _subtitulo_secao(document, "Leitura executiva por grupo")
    cedente = grupos_df[grupos_df["grupo"] == metricas.CEDENTE].iloc[0]
    destinataria = grupos_df[grupos_df["grupo"] == metricas.DESTINATARIA].iloc[0]
    neutra = grupos_df[grupos_df["grupo"] == metricas.NEUTRA].iloc[0]

    tabela = document.add_table(rows=1, cols=3)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    _definir_bordas_tabela(tabela)
    textos = [
        ("CEDENTES", cedente, False),
        ("DESTINATÁRIAS", destinataria, False),
        ("NEUTRAS", neutra, True),
    ]
    for i, (rotulo, linha, eh_neutra) in enumerate(textos):
        celula = tabela.rows[0].cells[i]
        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _sombrear_celula(celula, "F3F6F8")
        p = celula.paragraphs[0]
        run = p.add_run(rotulo + "\n")
        run.bold = True
        run.font.name = FONTE_PADRAO
        run.font.color.rgb = _RGB_TITULO
        if eh_neutra:
            corpo = (
                f"{formatar_numero(linha['total_com'])} processos\n"
                f"Média por unidade: {formatar_numero(linha['media_unidade_com'], 2)}\n"
                f"Média por magistrado: {formatar_numero(linha['media_magistrado_com'], 2)}\n"
                "Variação líquida: não se aplica"
            )
        else:
            corpo = (
                f"{formatar_numero(linha['total_sem'])} → {formatar_numero(linha['total_com'])} processos\n"
                f"Média por unidade: {formatar_numero(linha['media_unidade_sem'], 2)} → {formatar_numero(linha['media_unidade_com'], 2)}\n"
                f"Média por magistrado: {formatar_numero(linha['media_magistrado_sem'], 2)} → {formatar_numero(linha['media_magistrado_com'], 2)}\n"
                f"Variação líquida: {formatar_numero(linha['variacao_processos'])} processos "
                f"({formatar_percentual_com_sinal(linha['variacao_pct'], 1)})"
            )
        run2 = p.add_run(corpo)
        run2.font.name = FONTE_PADRAO
        run2.font.size = Pt(9.5)
    document.add_paragraph()

    _subtitulo_secao(document, "Quadro 2 – Convergência das médias entre os grupos")
    _paragrafo(document, "Diferença entre as médias das unidades cedentes e destinatárias", negrito=True, tamanho=10.5)
    _tabela_simples(document, ["Dimensão", "Sem equalização", "Com equalização", "Redução absoluta", "Redução relativa"], [
        ["Média de processos por unidade", formatar_numero(convergencia["unidade"]["sem"], 2),
         formatar_numero(convergencia["unidade"]["com"], 2), formatar_numero(convergencia["unidade"]["reducao_absoluta"], 2),
         formatar_percentual(convergencia["unidade"]["reducao_relativa_pct"], 1)],
        ["Média de processos por magistrado", formatar_numero(convergencia["magistrado"]["sem"], 2),
         formatar_numero(convergencia["magistrado"]["com"], 2), formatar_numero(convergencia["magistrado"]["reducao_absoluta"], 2),
         formatar_percentual(convergencia["magistrado"]["reducao_relativa_pct"], 1)],
    ], larguras_cm=[5, 3, 3, 3, 3], colunas_quantidade={1, 2, 3, 4})

    _subtitulo_secao(document, "Síntese interpretativa")
    verbo = "superava" if cedente["media_unidade_sem"] >= destinataria["media_unidade_sem"] else "era inferior à"
    _paragrafo(
        document,
        "O quadro compara a distância entre as médias dos dois grupos diretamente envolvidos no fluxo de "
        "equalização. As unidades cedentes são aquelas que redistribuem processos, enquanto as destinatárias "
        "recebem processos pelo sistema."
    )
    _paragrafo(
        document,
        f"No cenário simulado sem equalização, a média das unidades cedentes {verbo} a média das "
        f"destinatárias em {formatar_numero(convergencia['unidade']['sem'], 2)} processos por unidade. Após "
        f"as redistribuições, essa diferença caiu para {formatar_numero(convergencia['unidade']['com'], 2)} "
        f"processos, o que representa uma redução de {formatar_percentual(convergencia['unidade']['reducao_relativa_pct'], 1)}. "
        "O resultado indica uma forte aproximação entre os volumes médios de processos dos dois grupos."
    )
    _paragrafo(
        document,
        "A análise da carga de trabalho por magistrado apresenta comportamento semelhante. Sem a "
        f"equalização, a diferença entre as médias dos grupos seria de {formatar_numero(convergencia['magistrado']['sem'], 2)} "
        "processos por magistrado. No cenário com equalização, essa distância foi reduzida para "
        f"{formatar_numero(convergencia['magistrado']['com'], 2)} processos por magistrado, correspondendo a "
        f"uma redução de {formatar_percentual(convergencia['magistrado']['reducao_relativa_pct'], 1)}."
    )
    _paragrafo(
        document,
        "Os resultados não significam que todas as unidades passaram a possuir exatamente a mesma quantidade "
        "de processos ou a mesma carga de trabalho. A comparação demonstra que, em termos médios, a diferença "
        "estrutural entre unidades cedentes e destinatárias foi substancialmente reduzida. A medida por "
        "magistrado é especialmente relevante porque considera as diferenças de capacidade de trabalho entre "
        "as unidades, evitando que unidades com números distintos de magistrados sejam comparadas apenas pelo "
        "volume bruto de processos."
    )
    _paragrafo(
        document,
        "As unidades neutras não integram este quadro porque não enviam nem recebem processos pelo fluxo de "
        "equalização. Sua situação é considerada nas análises gerais de dispersão, mas não na comparação "
        "direta entre os dois grupos alcançados pelas redistribuições."
    )
