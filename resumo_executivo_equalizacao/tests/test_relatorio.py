"""Teste de fumaca fim-a-fim: gera o DOCX a partir de um consolidado
fixture e valida a estrutura esperada (secoes, orientacao, tabelas,
imagens, titulos de secao)."""
import zipfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

from src.relatorio import gerar_relatorio
from tests.helpers import escrever_consolidado_fixture

_TITULOS_ESPERADOS = [
    "1. Critérios de leitura",
    "2. Resultado geral do período",
    "3. Dimensão 1 — quantidade de processos por unidade",
    "4. Dimensão 2 — carga de trabalho por magistrado",
    "5. Comparação entre cedentes, destinatárias e neutras",
    "Anexo 1. Unidades cedentes",
    "Anexo 2. Unidades destinatárias",
    "Anexo 3 – Resumo por Grupo de Unidade",
]


def test_gera_docx_com_estrutura_esperada(tmp_path):
    caminho_consolidado = escrever_consolidado_fixture(tmp_path / "consolidado.xlsx")
    caminho_saida = tmp_path / "resumo.docx"

    resultado = gerar_relatorio(caminho_consolidado, caminho_saida)
    assert resultado == caminho_saida
    assert caminho_saida.exists()

    d = Document(str(caminho_saida))

    # duas secoes: corpo em retrato (Carta), anexos em paisagem
    assert len(d.sections) == 2
    assert d.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert d.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert round(d.sections[0].page_width.cm, 1) == 21.6
    assert round(d.sections[1].page_width.cm, 1) == 27.9

    titulos = [p.text for p in d.paragraphs for r in p.runs
               if r.bold and r.font.size == Pt(18)]
    for titulo in _TITULOS_ESPERADOS:
        assert titulo in titulos

    assert len(d.tables) >= 14

    # rodape com campo de paginacao
    rodape_xml = d.sections[0].footer.paragraphs[0]._p.xml
    assert "PAGE" in rodape_xml and "NUMPAGES" in rodape_xml


def test_docx_contem_graficos_e_logo(tmp_path):
    caminho_consolidado = escrever_consolidado_fixture(tmp_path / "consolidado.xlsx")
    caminho_saida = tmp_path / "resumo.docx"
    gerar_relatorio(caminho_consolidado, caminho_saida)

    with zipfile.ZipFile(caminho_saida) as z:
        imagens = [n for n in z.namelist() if n.startswith("word/media/")]
    # logo (reaproveitado nos 2 cabecalhos) + 3 graficos (Figura 1, 2 e 3)
    assert len(imagens) >= 4


def test_gera_docx_com_periodo_explicito(tmp_path):
    import pandas as pd
    caminho_consolidado = escrever_consolidado_fixture(tmp_path / "consolidado.xlsx")
    caminho_saida = tmp_path / "resumo_periodo.docx"

    gerar_relatorio(
        caminho_consolidado, caminho_saida,
        inicio=pd.Timestamp("2026-07-01"), fim=pd.Timestamp("2026-07-31"),
    )
    assert caminho_saida.exists()
