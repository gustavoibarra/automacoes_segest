"""Geracao do relatorio gerencial em DOCX (secao 20 do prompt), a partir
dos filtros selecionados. Utiliza python-docx; os graficos sao gerados em
alta resolucao a partir das mesmas figuras Plotly do dashboard (via
kaleido), com contingencia em matplotlib caso o kaleido nao esteja
disponivel no ambiente.
"""
from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src import charts
from src.formatting import formatar_data, formatar_numero, formatar_percentual

A4_LARGURA = Cm(21)
A4_ALTURA = Cm(29.7)
COR_TITULO = RGBColor(0x1F, 0x3A, 0x5F)


def _config_pagina_a4(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = A4_LARGURA
        section.page_height = A4_ALTURA
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def _marcar_linha_cabecalho_repetido(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _adicionar_campo_pagina(paragraph, texto_antes: str = "", texto_depois: str = "") -> None:
    run = paragraph.add_run(texto_antes)

    def _campo(tipo: str):
        r = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {tipo} "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_end)

    _campo("PAGE")
    paragraph.add_run(" de ")
    _campo("NUMPAGES")
    paragraph.add_run(texto_depois)


def _rodape(document: Document, data_atualizacao: datetime) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        _adicionar_campo_pagina(paragraph, texto_antes="Pagina ",
                                 texto_depois=f" | Atualizado em {data_atualizacao.strftime('%d/%m/%Y %H:%M')}")


def _titulo(document: Document, texto: str, nivel: int = 1):
    h = document.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = COR_TITULO
    return h


def _tabela_de_dataframe(document: Document, df: pd.DataFrame, formatadores: dict | None = None,
                          largura_total_cm: float = 17.0):
    formatadores = formatadores or {}
    if df.empty:
        document.add_paragraph("Sem dados para os filtros selecionados.")
        return

    tabela = document.add_table(rows=1, cols=len(df.columns))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Light Grid Accent 1"

    largura_coluna = Cm(largura_total_cm / len(df.columns))
    for i, coluna in enumerate(df.columns):
        celula = tabela.rows[0].cells[i]
        celula.text = str(coluna)
        for p in celula.paragraphs:
            for r in p.runs:
                r.bold = True
        celula.width = largura_coluna
    _marcar_linha_cabecalho_repetido(tabela.rows[0])

    for _, linha in df.iterrows():
        celulas = tabela.add_row().cells
        for i, coluna in enumerate(df.columns):
            valor = linha[coluna]
            if coluna in formatadores:
                texto = formatadores[coluna](valor)
            else:
                texto = "" if pd.isna(valor) else str(valor)
            celulas[i].text = texto
            celulas[i].width = largura_coluna
    document.add_paragraph()


def _figura_para_imagem(fig, caminho_png: Path, largura_px: int = 1600, altura_px: int = 900) -> bool:
    try:
        fig.write_image(str(caminho_png), width=largura_px, height=altura_px, scale=2)
        return True
    except Exception:
        try:
            import matplotlib.pyplot as plt
            plt.figure(figsize=(10, 5.6))
            plt.title(fig.layout.title.text if fig.layout.title else "")
            for trace in fig.data:
                try:
                    plt.bar(trace.x, trace.y, label=trace.name)
                except Exception:
                    continue
            plt.legend()
            plt.tight_layout()
            plt.savefig(caminho_png, dpi=200)
            plt.close()
            return True
        except Exception:
            return False


def _adicionar_grafico(document: Document, fig, tmp_dir: Path, nome_arquivo: str, largura_cm: float = 16.0) -> None:
    caminho = tmp_dir / nome_arquivo
    if _figura_para_imagem(fig, caminho):
        document.add_picture(str(caminho), width=Cm(largura_cm))
    else:
        document.add_paragraph("[Nao foi possivel renderizar o grafico neste ambiente]")


def gerar_relatorio_docx(caminho_saida: Path, *, filtros, metricas: dict, processos_filtrados: pd.DataFrame,
                          unidades_filtradas: pd.DataFrame, arquivos_utilizados: list[str],
                          qualidade: dict, anomalias_df: pd.DataFrame, periodo_texto: str,
                          filtros_texto: str, data_geracao: datetime | None = None) -> Path:
    data_geracao = data_geracao or datetime.now()

    document = Document()
    _config_pagina_a4(document)

    # ---- Capa -------------------------------------------------------
    document.add_paragraph()
    document.add_paragraph()
    capa_titulo = document.add_paragraph()
    capa_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa_titulo.add_run("Relatorio Gerencial - Sistema de Equalizacao TRT12")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = COR_TITULO

    subtitulo = document.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run("Analise de distribuicao e redistribuicao de processos judiciais").italic = True

    document.add_paragraph()
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Periodo analisado: {periodo_texto}\n")
    info.add_run(f"Data de geracao: {formatar_data(data_geracao)} {data_geracao.strftime('%H:%M')}\n")
    info.add_run(f"Arquivos utilizados: {len(arquivos_utilizados)}")
    document.add_page_break()

    # ---- Arquivos utilizados e filtros -------------------------------
    _titulo(document, "1. Arquivos utilizados e filtros aplicados")
    document.add_paragraph("Arquivos CSV consolidados nesta analise:")
    for nome in arquivos_utilizados:
        document.add_paragraph(nome, style="List Bullet")
    document.add_paragraph(f"Filtros aplicados: {filtros_texto or 'nenhum filtro adicional'}")
    document.add_page_break()

    # ---- Resumo gerencial / cards -------------------------------------
    _titulo(document, "2. Resumo gerencial e cards executivos")
    cards = metricas["cards"]
    resumo = (
        f"No periodo analisado foram identificados {formatar_numero(cards['casos_novos'])} casos novos, "
        f"dos quais {formatar_numero(cards['processos_equalizados'])} "
        f"({formatar_percentual(cards['percentual_processos_equalizados'])}) foram equalizados validamente pelo "
        f"Sistema de Equalizacao. A adocao do sistema reduziu a dispersao entre unidades em "
        f"{formatar_percentual(cards['reducao_dispersao_pct'])} em relacao ao cenario simulado sem equalizacao."
    )
    document.add_paragraph(resumo)

    cards_df = pd.DataFrame([
        {"Indicador": "Casos novos (COUNT DISTINCT processo)", "Valor": formatar_numero(cards["casos_novos"])},
        {"Indicador": "Processos equalizados (COUNT DISTINCT processo)", "Valor": formatar_numero(cards["processos_equalizados"])},
        {"Indicador": "% processos equalizados", "Valor": formatar_percentual(cards["percentual_processos_equalizados"])},
        {"Indicador": "Media equalizados por unidade destinataria", "Valor": formatar_numero(cards["media_equalizados_por_unidade_destinataria"], 2)},
        {"Indicador": "Media de processos por unidade cedente", "Valor": formatar_numero(cards["media_processos_por_unidade_cedente"], 2)},
        {"Indicador": "Media de processos por unidade neutra", "Valor": formatar_numero(cards["media_processos_por_unidade_neutra"], 2)},
        {"Indicador": "Media de processos por unidade destinataria", "Valor": formatar_numero(cards["media_processos_por_unidade_destinataria"], 2)},
        {"Indicador": "Reducao da dispersao (MAD)", "Valor": formatar_percentual(cards["reducao_dispersao_pct"])},
        {"Indicador": "Media ponderada por magistrado (com equalizacao)",
         "Valor": formatar_numero(cards["media_ponderada_magistrado_com_equalizacao"], 2)},
        {"Indicador": "Media ponderada por magistrado (sem equalizacao)",
         "Valor": formatar_numero(cards["media_ponderada_magistrado_sem_equalizacao"], 2)},
        {"Indicador": "Reducao da dispersao por magistrado (MAD)",
         "Valor": formatar_percentual(cards["reducao_dispersao_magistrado_pct"])},
    ])
    _tabela_de_dataframe(document, cards_df)
    document.add_page_break()

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)

        # ---- Comparativo com e sem equalizacao -------------------------
        _titulo(document, "3. Comparativo com e sem Sistema de Equalizacao")
        fig_comp = charts.grafico_comparativo_barras(metricas["tabela_com"], metricas["tabela_sem"],
                                                       periodo_texto=periodo_texto, filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_comp, tmp_dir, "comparativo.png")

        fig_disp = charts.grafico_dispersao_45graus(metricas["tabela_com"], metricas["tabela_sem"],
                                                      periodo_texto=periodo_texto, filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_disp, tmp_dir, "dispersao.png")

        fig_box = charts.boxplot_por_grupo(metricas["tabela_com"], metricas["tabela_sem"],
                                            periodo_texto=periodo_texto, filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_box, tmp_dir, "boxplot.png")

        disp_mag_com_preview, disp_mag_sem_preview = metricas["dispersao_magistrado_com"], metricas["dispersao_magistrado_sem"]
        magistrado_disponivel = not (pd.isna(disp_mag_com_preview["media"]) and pd.isna(disp_mag_sem_preview["media"]))
        if magistrado_disponivel:
            _titulo(document, "3.1 Comparativo por magistrado", nivel=2)
            fig_comp_mag = charts.grafico_comparativo_barras_por_magistrado(
                metricas["tabela_com"], metricas["tabela_sem"], periodo_texto=periodo_texto, filtros_texto=filtros_texto)
            _adicionar_grafico(document, fig_comp_mag, tmp_dir, "comparativo_magistrado.png")

            fig_disp_mag = charts.grafico_dispersao_45graus_por_magistrado(
                metricas["tabela_com"], metricas["tabela_sem"], periodo_texto=periodo_texto, filtros_texto=filtros_texto)
            _adicionar_grafico(document, fig_disp_mag, tmp_dir, "dispersao_magistrado.png")

            fig_box_mag = charts.boxplot_por_grupo_por_magistrado(
                metricas["tabela_com"], metricas["tabela_sem"], periodo_texto=periodo_texto, filtros_texto=filtros_texto)
            _adicionar_grafico(document, fig_box_mag, tmp_dir, "boxplot_magistrado.png")
        document.add_page_break()

        # ---- Indicadores de dispersao -----------------------------------
        _titulo(document, "4. Indicadores de dispersao")
        disp_com, disp_sem = metricas["dispersao_com"], metricas["dispersao_sem"]
        comp = metricas["comparacao_dispersao"]
        disp_df = pd.DataFrame([
            {"Medida": "Media", "Sem equalizacao": disp_sem["media"], "Com equalizacao": disp_com["media"]},
            {"Medida": "MAD (distancia media absoluta)", "Sem equalizacao": disp_sem["mad"], "Com equalizacao": disp_com["mad"]},
            {"Medida": "Desvio-padrao populacional", "Sem equalizacao": disp_sem["desvio_padrao"], "Com equalizacao": disp_com["desvio_padrao"]},
            {"Medida": "Coeficiente de variacao (%)", "Sem equalizacao": disp_sem["coeficiente_variacao"], "Com equalizacao": disp_com["coeficiente_variacao"]},
            {"Medida": "Amplitude", "Sem equalizacao": disp_sem["amplitude"], "Com equalizacao": disp_com["amplitude"]},
            {"Medida": "Unidades acima da media", "Sem equalizacao": disp_sem["qtd_acima_media"], "Com equalizacao": disp_com["qtd_acima_media"]},
            {"Medida": "Unidades abaixo da media", "Sem equalizacao": disp_sem["qtd_abaixo_media"], "Com equalizacao": disp_com["qtd_abaixo_media"]},
        ])
        _tabela_de_dataframe(document, disp_df, formatadores={
            "Sem equalizacao": lambda v: formatar_numero(v, 2), "Com equalizacao": lambda v: formatar_numero(v, 2),
        })
        document.add_paragraph(
            f"Reducao da MAD: {formatar_percentual(comp['reducao_mad_pct'])} | "
            f"Reducao do desvio-padrao: {formatar_percentual(comp['reducao_desvio_padrao_pct'])} | "
            f"Reducao do coeficiente de variacao: {formatar_percentual(comp['reducao_coeficiente_variacao_pct'])} | "
            f"Situacao: {'melhora' if comp['melhora'] else 'piora'}."
        )

        if magistrado_disponivel:
            _titulo(document, "4.1 Indicadores de dispersao por magistrado", nivel=2)
            comp_mag = metricas["comparacao_dispersao_magistrado"]
            disp_mag_df = pd.DataFrame([
                {"Medida": "Media ponderada (processos/magistrado)", "Sem equalizacao": disp_mag_sem_preview["media"], "Com equalizacao": disp_mag_com_preview["media"]},
                {"Medida": "MAD (distancia media absoluta)", "Sem equalizacao": disp_mag_sem_preview["mad"], "Com equalizacao": disp_mag_com_preview["mad"]},
                {"Medida": "Desvio-padrao (em torno da media ponderada)", "Sem equalizacao": disp_mag_sem_preview["desvio_padrao"], "Com equalizacao": disp_mag_com_preview["desvio_padrao"]},
                {"Medida": "Coeficiente de variacao (%)", "Sem equalizacao": disp_mag_sem_preview["coeficiente_variacao"], "Com equalizacao": disp_mag_com_preview["coeficiente_variacao"]},
                {"Medida": "Amplitude", "Sem equalizacao": disp_mag_sem_preview["amplitude"], "Com equalizacao": disp_mag_com_preview["amplitude"]},
                {"Medida": "Unidades acima da media", "Sem equalizacao": disp_mag_sem_preview["qtd_acima_media"], "Com equalizacao": disp_mag_com_preview["qtd_acima_media"]},
                {"Medida": "Unidades abaixo da media", "Sem equalizacao": disp_mag_sem_preview["qtd_abaixo_media"], "Com equalizacao": disp_mag_com_preview["qtd_abaixo_media"]},
            ])
            _tabela_de_dataframe(document, disp_mag_df, formatadores={
                "Sem equalizacao": lambda v: formatar_numero(v, 2), "Com equalizacao": lambda v: formatar_numero(v, 2),
            })
            document.add_paragraph(
                "Media ponderada pela quantidade de magistrados: total de processos / total de magistrados das "
                "unidades com quantidade de magistrados informada (nao e a media simples das taxas por unidade)."
            )
            document.add_paragraph(
                f"Reducao da MAD por magistrado: {formatar_percentual(comp_mag['reducao_mad_pct'])} | "
                f"Reducao do desvio-padrao: {formatar_percentual(comp_mag['reducao_desvio_padrao_pct'])} | "
                f"Reducao do coeficiente de variacao: {formatar_percentual(comp_mag['reducao_coeficiente_variacao_pct'])} | "
                f"Situacao: {'melhora' if comp_mag['melhora'] else 'piora'}."
            )
        document.add_page_break()

        # ---- Analise por grupo ------------------------------------------
        _titulo(document, "5. Analise por grupo de unidades")
        grupo_df = metricas["analise_grupo"].copy()
        _tabela_de_dataframe(document, grupo_df, formatadores={
            c: (lambda v: formatar_numero(v, 2)) for c in
            ["media_com_equalizacao", "media_sem_equalizacao", "mad_com_equalizacao", "mad_sem_equalizacao",
             "variacao_media_pct", "variacao_dispersao_pct", "media_ponderada_magistrado_com_equalizacao",
             "media_ponderada_magistrado_sem_equalizacao", "mad_ponderada_magistrado_com_equalizacao",
             "mad_ponderada_magistrado_sem_equalizacao", "variacao_media_magistrado_pct",
             "variacao_dispersao_magistrado_pct"]
        })
        document.add_page_break()

        # ---- Processos cedidos / destinados / fluxos ---------------------
        _titulo(document, "6. Processos cedidos por unidade")
        _tabela_de_dataframe(document, metricas["cedidos"].drop(columns=["unidade_norm"], errors="ignore"),
                              formatadores={"percentual_do_total_cedido": formatar_percentual,
                                            "media_diaria": lambda v: formatar_numero(v, 2),
                                            "media_semanal": lambda v: formatar_numero(v, 2)})
        fig_cedidos = charts.grafico_barras_unidade(metricas["cedidos"], "quantidade_processos_cedidos", "unidade",
                                                     "Processos cedidos por unidade", periodo_texto=periodo_texto,
                                                     filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_cedidos, tmp_dir, "cedidos.png")
        document.add_page_break()

        _titulo(document, "7. Processos destinados por unidade")
        _tabela_de_dataframe(document, metricas["destinados"].drop(columns=["unidade_norm"], errors="ignore"),
                              formatadores={"percentual_do_total_recebido": formatar_percentual,
                                            "media_diaria": lambda v: formatar_numero(v, 2),
                                            "media_semanal": lambda v: formatar_numero(v, 2)})
        fig_destinados = charts.grafico_barras_unidade(metricas["destinados"], "quantidade_processos_recebidos", "unidade",
                                                         "Processos recebidos por unidade", periodo_texto=periodo_texto,
                                                         filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_destinados, tmp_dir, "destinados.png")
        document.add_page_break()

        _titulo(document, "8. Principais fluxos de equalizacao")
        fluxos = metricas["fluxos"]
        document.add_paragraph(
            f"Quantidade de pares cedente-destinataria: {formatar_numero(fluxos['quantidade_pares'])} | "
            f"Participacao das dez maiores rotas: {formatar_percentual(fluxos['participacao_top10_pct'])}"
        )
        if not fluxos["rotas"].empty:
            _tabela_de_dataframe(document, fluxos["principais_rotas"], formatadores={
                "participacao_pct": formatar_percentual,
            })
            fig_sankey = charts.sankey_fluxos(fluxos["rotas"], periodo_texto=periodo_texto, filtros_texto=filtros_texto)
            _adicionar_grafico(document, fig_sankey, tmp_dir, "sankey.png")
        else:
            document.add_paragraph("Nenhum episodio de equalizacao valida no filtro atual.")
        document.add_page_break()

        # ---- Evolucao temporal ---------------------------------------
        _titulo(document, "9. Evolucao temporal")
        adicionais = metricas["adicionais"]
        fig_evol = charts.grafico_evolucao(adicionais["evolucao_diaria_casos_novos"],
                                            adicionais["evolucao_semanal_casos_novos"],
                                            "Evolucao de casos novos", periodo_texto=periodo_texto,
                                            filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_evol, tmp_dir, "evolucao_casos.png")
        fig_evol_eq = charts.grafico_evolucao(adicionais["evolucao_diaria_equalizacoes"],
                                               adicionais["evolucao_semanal_equalizacoes"],
                                               "Evolucao de equalizacoes validas", periodo_texto=periodo_texto,
                                               filtros_texto=filtros_texto)
        _adicionar_grafico(document, fig_evol_eq, tmp_dir, "evolucao_eq.png")
        document.add_paragraph(
            f"Tempo medio entre distribuicao e destino da equalizacao: "
            f"{formatar_numero(adicionais['tempo_medio_distribuicao_destino_dias'], 1)} dias "
            f"(mediana: {formatar_numero(adicionais['tempo_mediano_distribuicao_destino_dias'], 1)} dias). "
            f"Tempo medio de permanencia na triagem: "
            f"{formatar_numero(adicionais['tempo_medio_permanencia_triagem_dias'], 1)} dias "
            f"(mediana: {formatar_numero(adicionais['tempo_mediano_permanencia_triagem_dias'], 1)} dias)."
        )
        document.add_page_break()

        # ---- Classes judiciais ------------------------------------------
        _titulo(document, "10. Principais classes judiciais")
        classes_df = adicionais["percentual_equalizados_por_classe"].sort_values("casos_novos", ascending=False).head(15)
        _tabela_de_dataframe(document, classes_df, formatadores={"percentual_equalizados": formatar_percentual})
        document.add_page_break()

        # ---- Anomalias e qualidade ----------------------------------------
        _titulo(document, "11. Anomalias e qualidade dos dados")
        qual_df = pd.DataFrame([
            {"Indicador": "Registros lidos", "Valor": formatar_numero(qualidade["quantidade_registros_lidos"])},
            {"Indicador": "Quantidade de arquivos", "Valor": formatar_numero(qualidade["quantidade_arquivos"])},
            {"Indicador": "Duplicidades removidas", "Valor": formatar_numero(qualidade["duplicidades_removidas"])},
            {"Indicador": "Registros validos", "Valor": formatar_numero(qualidade["registros_validos"])},
            {"Indicador": "Registros com anomalia", "Valor": formatar_numero(qualidade["registros_com_anomalia"])},
            {"Indicador": "Registros RESTITUIDO (desconsiderados de todas as analises)", "Valor": formatar_numero(qualidade.get("registros_restituidos", 0))},
            {"Indicador": "Processos distintos", "Valor": formatar_numero(qualidade["processos_distintos"])},
            {"Indicador": "Unidades nao classificadas", "Valor": formatar_numero(qualidade["unidades_nao_classificadas"])},
            {"Indicador": "Data minima", "Valor": formatar_data(qualidade["data_minima"])},
            {"Indicador": "Data maxima", "Valor": formatar_data(qualidade["data_maxima"])},
        ])
        _tabela_de_dataframe(document, qual_df)

        if not anomalias_df.empty:
            resumo_anomalias = anomalias_df.groupby("tipo_anomalia").size().reset_index(name="quantidade")
            resumo_anomalias = resumo_anomalias.sort_values("quantidade", ascending=False)
            document.add_paragraph("Anomalias por tipo:")
            _tabela_de_dataframe(document, resumo_anomalias)
        document.add_page_break()

        # ---- Metodologia --------------------------------------------------
        _titulo(document, "12. Metodologia e formulas")
        document.add_paragraph(_TEXTO_METODOLOGIA)

    _rodape(document, data_geracao)
    caminho_saida.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(caminho_saida))
    return caminho_saida


_TEXTO_METODOLOGIA = """
Casos novos: quantidade de numeros de processo distintos no periodo e filtros selecionados (COUNT DISTINCT processo).
Processos equalizados: quantidade de processos distintos com equalizacao valida (COUNT DISTINCT processo WHERE equalizacao_valida = verdadeiro).
Uma equalizacao e detectada quando o historico cronologico do processo contem a sequencia UNIDADE PERMANENTE -> UNIDADE DE TRIAGEM -> UNIDADE PERMANENTE. E considerada valida quando a unidade anterior a triagem e classificada como CEDENTE, a primeira unidade permanente posterior a triagem e classificada como DESTINATARIA, e os movimentos de entrada na triagem e no destino sao redistribuicoes.
Cenario com equalizacao: cada processo e contado na sua unidade final real (ultima unidade permanente do historico).
Cenario sem equalizacao (simulado): processos com equalizacao valida sao recontados na unidade cedente imediatamente anterior a triagem; os demais permanecem na unidade final real.
MAD (distancia media absoluta): media do modulo da diferenca entre a quantidade de cada unidade e a media geral do cenario.
Reducao da MAD: (MAD sem equalizacao - MAD com equalizacao) / MAD sem equalizacao x 100.
Todos os indicadores informam se a contagem e distinta por processo, por movimentacao, por episodio de equalizacao ou por unidade, conforme aplicavel.
""".strip()
